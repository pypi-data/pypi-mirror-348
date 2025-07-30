import os
import regex as re
import json
from dotenv import load_dotenv
load_dotenv()

from azure.core.exceptions import ResourceNotFoundError as AzureResourceNotFoundError

class ResourceNotFoundError(AzureResourceNotFoundError):
    def __init__(self, message: str):
        super().__init__(message)

from collections import defaultdict
from enum import Enum
from time import time 
from typing import Any, Callable, List, Dict, Optional, Tuple, Union, ClassVar
from datetime import datetime, timezone # Add datetime imports

from azure.identity import ClientSecretCredential, DefaultAzureCredential
from azure.core.credentials import AccessToken, TokenCredential
from azure.mgmt.resource import SubscriptionClient, ResourceManagementClient
from azure.mgmt.storage import StorageManagementClient
import azure.mgmt.storage.models as azstm
from azure.mgmt.storage.models import StorageAccount
import azure.mgmt.resource.subscriptions.models as azsbm
from azure.mgmt.search import SearchManagementClient
import azure.mgmt.search.models as azsrm
import azure.search.documents as azsd
from azure.search.documents.models import VectorizedQuery, VectorizableTextQuery, QueryType
from azure.search.documents.indexes.models import (
    SearchIndexer,
    IndexingParameters,
    FieldMapping,
    OutputFieldMappingEntry
)
from tenacity import retry, stop_after_attempt, wait_random_exponential
class Identity:
    """Azure Identity for authentication.

    Two authentication methods are supported, both implementing TokenCredential:
    1. Default authentication (when no arguments provided):
       - Uses DefaultAzureCredential for automatic authentication
    2. Service Principal authentication (when any argument provided):
       - Requires all three: tenant_id, client_id, client_secret
       - Uses ClientSecretCredential
    """

    credential: TokenCredential
    tenant_id: Optional[str] = None
    client_id: Optional[str] = None
    client_secret: Optional[str] = None

    subscription_client: SubscriptionClient

    def __init__(self, tenant_id: Optional[str] = None, client_id: Optional[str] = None, client_secret: Optional[str] = None):
        # Use DefaultAzureCredential when all parameters are None
        if tenant_id is None and client_id is None and client_secret is None:
            self.credential = DefaultAzureCredential()
        else:
            # Require all parameters for ClientSecretCredential
            if not all([tenant_id, client_id, client_secret]):
                raise ValueError("For client credential auth, all three parameters (tenant_id, client_id, client_secret) are required")

            self.tenant_id = tenant_id
            self.client_id = client_id
            self.client_secret = client_secret
            self.credential = ClientSecretCredential(tenant_id=tenant_id, client_id=client_id, client_secret=client_secret)

        # Validate credentials
        token:AccessToken = self.credential.get_token("https://management.azure.com/.default")
        if token is None:
            raise ValueError("Failed to get token. Check your credentials.")
        self.subscription_client = SubscriptionClient(self.credential)
    
    def get_credential(self) -> TokenCredential:
        return self.credential
    
    def get_subscriptions(self) -> List[azsbm.Subscription]:
        subscriptions = list(self.subscription_client.subscriptions.list())
        return subscriptions 
    
    def get_subscription(self, subscription_id: str) -> "Subscription":
        for sub in self.get_subscriptions():
            if sub.subscription_id == subscription_id:
                return Subscription(self, sub, sub.subscription_id)
        raise ValueError(f"Subscription with ID {subscription_id} not found.")

from azure.mgmt.cognitiveservices import CognitiveServicesManagementClient

class Subscription:

    identity: Identity
    subscription: azsbm.Subscription
    subscription_id: str

    resource_client: ResourceManagementClient
    storage_client: StorageManagementClient

    def __init__(self, identity: Identity, 
                 subscription: azsbm.Subscription, subscription_id: str):
        self.identity = identity
        self.subscription = subscription
        self.subscription_id = subscription_id
        self.resource_client = ResourceManagementClient(self.identity.get_credential(), self.subscription_id)
        self.storage_client = StorageManagementClient(self.identity.get_credential(), self.subscription_id)

    def get_resource_group(self, group_name: str) -> "ResourceGroup":
        groups = self.resource_client.resource_groups.list()
        for group in groups:
            if group.name.lower() == group_name.lower():
                return ResourceGroup(self, group)
        raise ValueError(f"Resource group with name {group_name} not found.")
    
    def create_resource_group(self, group_name: str, location: str) -> "ResourceGroup":
        result = self.resource_client.resource_groups.create_or_update(
            group_name,
            {"location": location}
        )
        if result is not None:
            return ResourceGroup(self, result)
        raise ValueError(f"Failed to create resource group with name {group_name}.")
    
    def get_search_services(self) -> List[azsrm.SearchService]:
        search_mgmt_client = SearchManagementClient(self.identity.get_credential(), self.subscription_id)
        services = list(search_mgmt_client.services.list_by_subscription())
        return services 
    
    def get_search_service(self, service_name: str) -> "SearchService":
        services = self.get_search_services()
        for service in services:
            if service.name == service_name:
                resource_group_name = service.id.split("/")[4] 
                resource_group = self.get_resource_group(resource_group_name)
                return SearchService(self, resource_group, service)
        raise ValueError(f"Search service with name {service_name} not found.")
    
    def get_storage_management_client(self) -> StorageManagementClient:
        if self.storage_client is None:
            self.storage_client = StorageManagementClient(self.identity.get_credential(), self.subscription_id) 
        return self.storage_client
    
    def get_storage_accounts(self) -> List[azstm.StorageAccount]:
        accounts = list(self.storage_client.storage_accounts.list())
        return accounts
    
    def get_cognitive_client(self) -> CognitiveServicesManagementClient: 
        cognitive_client: CognitiveServicesManagementClient 
        cognitive_client = CognitiveServicesManagementClient(
            credential=self.identity.get_credential(), 
            subscription_id=self.subscription_id
        )
        return cognitive_client
                
import azure.mgmt.resource.resources.models as azrm
import azure.mgmt.cognitiveservices.models as azcsm
class ResourceGroup:
    azure_resource_group: azrm.ResourceGroup
    subscription: Subscription

    def __init__(self, subscription: Subscription, azure_resource_group: azrm.ResourceGroup):
        self.subscription = subscription
        self.azure_resource_group = azure_resource_group

    def get_name(self) -> str:
        return self.azure_resource_group.name

    def get_resources(self) -> List[azrm.GenericResource]:
        resources = self.subscription.resource_client.resources.list_by_resource_group(self.azure_resource_group.name)
        return resources
    
    def get_storage_management_client(self) -> StorageManagementClient:
        return self.subscription.get_storage_management_client()
    
    def create_search_service(self, name: str, location: str) -> "SearchService":
        search_mgmt_client = SearchManagementClient(self.subscription.identity.get_credential(), 
                                                    self.subscription.subscription_id)
        # Define the search service
        search_service: Dict[str, Any] = {
            "location": location,
            "sku": {
                "name": "basic"  # Options: free, basic, standard, standard2, standard3, storage_optimized_l1, storage_optimized_l2
            }
        }
        operation = search_mgmt_client.services.begin_create_or_update(
            resource_group_name=self.azure_resource_group.name,
            search_service_name=name,
            service=search_service
        )
        search_service = operation.result()
        return SearchService(self, search_service)
    
    def get_storage_account(self, account_name: str) -> Optional["StorageAccount"]:
        storage_client = self.subscription.get_storage_management_client()
        try:
            account = storage_client.storage_accounts.get_properties(resource_group_name = self.azure_resource_group.name, account_name = account_name)
        except AzureResourceNotFoundError as e:
            raise ResourceNotFoundError(f"Storage account with name {account_name} not found. \n {str(e)}")
        except Exception as e:
            print(f"Error at ResourceGroup.get_storage_account({account_name = }): {str(e)}")
            raise e
        if account is None:
            raise ValueError(f"Storage account with name {account_name} not found.")
        return StorageAccount(self, account)

    def create_storage_account(self, account_name: str, location: str) -> "StorageAccount":
        storage_client = self.subscription.get_storage_management_client()
        params = azstm.StorageAccountCreateParameters(
            sku=azstm.Sku(name="Standard_LRS"), 
            kind=azstm.Kind.STORAGE_V2, 
            location=location
        )
        result = storage_client.storage_accounts.begin_create(resource_group_name=self.azure_resource_group.name, 
                                                              account_name=account_name, 
                                                              parameters=params)
        account = result.result()
        return StorageAccount(self, account)

    def get_ai_service(self, service_name:str) -> Optional["AIService"]:
        cognitive_client = self.subscription.get_cognitive_client()
        cognitive_accounts = cognitive_client.accounts.list_by_resource_group(self.azure_resource_group.name)

        accounts: List[azcsm.Account] = [account for account in cognitive_accounts]
        openai_services_names: List[str] = [account.name for account in accounts]

        for account in accounts :
            if account.kind.lower() == "openai" and account.name.lower() == service_name.lower() :
                return AIService(self, cognitive_client=cognitive_client, azure_Account=account)
        return None
        
from azure.storage.blob import BlobServiceClient, ContainerProperties, ContainerClient, BlobProperties
def validate_container_name(container_name: str) -> bool:
        # Validate container name
        if not container_name or not container_name.islower() or len(container_name) < 3 or len(container_name) > 63:
            raise ValueError("Container name must be lowercase, between 3-63 characters")
            
        if not all(c.islower() or c.isdigit() or c == '-' for c in container_name):
            raise ValueError("Container name can only contain lowercase letters, numbers, and hyphens")
            
        if container_name[0] == '-' or container_name[-1] == '-' or '--' in container_name:
            raise ValueError("Container name cannot start or end with hyphens or contain consecutive hyphens")
        return True
from azure.storage.blob import PublicAccess

class StorageAccount: 
    storage_account: azstm.StorageAccount
    resource_group: ResourceGroup

    storage_key: str
    connection_string_description: str

    def __init__(self, resource_group: ResourceGroup, storage_account: azstm.StorageAccount):
        self.resource_group = resource_group
        self.storage_account = storage_account
        client = resource_group.get_storage_management_client()
        keys = client.storage_accounts.list_keys(resource_group_name=resource_group.get_name(), 
                                                 account_name=storage_account.name)
        self.storage_key = keys.keys[0].value
        self.connection_string_description = f"DefaultEndpointsProtocol=https;AccountName={storage_account.name};AccountKey={self.storage_key};EndpointSuffix=core.windows.net"
    
    def get_name(self) -> str:
        return self.storage_account.name
    
    def get_blob_service_client(self) -> BlobServiceClient:
        return BlobServiceClient.from_connection_string(self.connection_string_description) 
    
    def get_container_client(self, container_name: str) -> ContainerClient:
        client = self.get_blob_service_client()
        container_client = client.get_container_client(container_name)
        return container_client

    def get_containers(self) -> List[ContainerProperties]:
        client = self.get_blob_service_client()
        containers = client.list_containers()
        return [container for container in containers]

    def get_container(self, container_name:str) -> "Container":
        client = self.get_blob_service_client()
        container = client.get_container_client(container_name)
        if container is None:
            raise ValueError(f"Container with name {container_name} not found.")
        return Container(self, container)
    
    def get_access_level(public_access_level: Optional[str] = None ) -> PublicAccess | None:
        # Convert public_access_level to appropriate enum value if provided
        access_level = None
        if public_access_level:
            if public_access_level.lower() == 'container':
                access_level = PublicAccess.Container
            elif public_access_level.lower() == 'blob':
                access_level = PublicAccess.Blob
            else:
                raise ValueError("public_access_level must be 'container', 'blob', or None")

    def create_container(self, container_name: str, public_access_level: Optional[str] = None) -> "Container":
        """
            Create a new container in the storage account.
            
            Args:
                container_name: Name of the container to create. Must be a valid DNS name,
                                lowercase, between 3-63 characters, and can contain only letters,
                                numbers, and hyphens.
                public_access_level: Optional level of public access. Valid values are:
                                    'container' (Full public read access for container and blob data),
                                    'blob' (Public read access for blobs only),
                                    None (No public access, default)
            
            Returns:
                Container: The created container object
                
            Raises:
                ValueError: If container name is invalid or if creation fails
                azure.core.exceptions.ResourceExistsError: If container already exists
        """
        validate_container_name(container_name)

        # Convert public_access_level to appropriate enum value if provided
        access_level = StorageAccount.get_access_level(public_access_level)

        
        # Get the blob service client
        blob_service_client = self.get_blob_service_client()
        
        # Create the container
        container_client = blob_service_client.create_container(
            name=container_name,
            public_access=access_level
        )
        
        print(f"Container '{container_name}' created successfully")
        
        # Return a Container object wrapping the new container
        return Container(self, container_client)

    def delete_container(self, container_name: str, force: bool = False) -> bool:
        """
        Delete a container from the storage account.
        
        Args:
            container_name: Name of the container to delete
            force: If True, deletes the container even if it contains blobs. 
                If False, will raise an error if the container is not empty.
        
        Returns:
            bool: True if the container was successfully deleted, False if it didn't exist
            
        Raises:
            ValueError: If deletion fails for reasons other than the container not existing
                    or if the container is not empty and force=False
        """
        # Get the blob service client
        blob_service_client = self.get_blob_service_client()
        
        # Check if the container exists
        container_client = blob_service_client.get_container_client(container_name)
        
        # Check if the container is empty when force is False
        if not force:
            # List blobs with maxresults=1 to check if any blobs exist
            blob_iterator = container_client.list_blobs()
            try:
                next(iter(blob_iterator))
                raise ValueError( f"Container '{container_name}' is not empty. Set force=True to delete anyway." )
            except StopIteration:
                # No blobs found, safe to delete
                pass    
        
        # Delete the container
        try:
            container_client.delete_container()
        except AzureResourceNotFoundError as e:
            raise ResourceNotFoundError(f"Container with name {container_name} not found.")
        print(f"Container '{container_name}' deleted successfully")
        return True

class BlobType(Enum):
    """
    Enumeration of common MIME types for blobs in Azure Storage
    """
    # Text formats
    TEXT_PLAIN = "text/plain"
    TEXT_CSV = "text/csv"
    TEXT_HTML = "text/html"
    TEXT_CSS = "text/css"
    TEXT_JAVASCRIPT = "text/javascript"
    TEXT_XML = "text/xml"
    TEXT_MARKDOWN = "text/markdown"
    
    # Application formats
    APP_JSON = "application/json"
    APP_XML = "application/xml"
    APP_PDF = "application/pdf"
    APP_ZIP = "application/zip"
    APP_GZIP = "application/gzip"
    APP_OCTET_STREAM = "application/octet-stream"
    
    # Microsoft Office formats
    MS_WORD = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    MS_EXCEL = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    MS_POWERPOINT = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    
    # Image formats
    IMAGE_JPEG = "image/jpeg"
    IMAGE_PNG = "image/png"
    IMAGE_GIF = "image/gif"
    IMAGE_SVG = "image/svg+xml"
    IMAGE_WEBP = "image/webp"
    IMAGE_TIFF = "image/tiff"
    
    # Audio formats
    AUDIO_MP3 = "audio/mpeg"
    AUDIO_WAV = "audio/wav"
    AUDIO_OGG = "audio/ogg"
    
    # Video formats
    VIDEO_MP4 = "video/mp4"
    VIDEO_WEBM = "video/webm"
    VIDEO_OGG = "video/ogg"
    
    

    # Common file extensions to MIME type mapping
    @classmethod
    def from_extension(cls, extension: str) -> Optional["BlobType"]:
        """
        Get MIME type from file extension
        
        Args:
            extension: The file extension (with or without leading dot)
        
        Returns:
            BlobType enum value or None if unknown extension
        """
        if not extension.startswith('.'):
            extension = '.' + extension
            
        extension = extension.lower()
        
        ext_map = {
            '.txt': cls.TEXT_PLAIN,
            '.csv': cls.TEXT_CSV,
            '.html': cls.TEXT_HTML,
            '.htm': cls.TEXT_HTML,
            '.css': cls.TEXT_CSS,
            '.js': cls.TEXT_JAVASCRIPT,
            '.xml': cls.TEXT_XML,
            '.md': cls.TEXT_MARKDOWN,
            '.json': cls.APP_JSON,
            '.pdf': cls.APP_PDF,
            '.zip': cls.APP_ZIP,
            '.gz': cls.APP_GZIP,
            '.docx': cls.MS_WORD,
            '.xlsx': cls.MS_EXCEL,
            '.pptx': cls.MS_POWERPOINT,
            '.jpg': cls.IMAGE_JPEG,
            '.jpeg': cls.IMAGE_JPEG,
            '.png': cls.IMAGE_PNG,
            '.gif': cls.IMAGE_GIF,
            '.svg': cls.IMAGE_SVG,
            '.webp': cls.IMAGE_WEBP,
            '.tif': cls.IMAGE_TIFF,
            '.tiff': cls.IMAGE_TIFF,
            '.mp3': cls.AUDIO_MP3,
            '.wav': cls.AUDIO_WAV,
            '.ogg': cls.AUDIO_OGG,
            '.mp4': cls.VIDEO_MP4,
            '.webm': cls.VIDEO_WEBM,
            '.bin': cls.APP_OCTET_STREAM
        }
        
        return ext_map.get(extension)
        
    @classmethod
    def from_mime_type(cls, mime_type: str) -> Optional["BlobType"]:
        """
        Get BlobType from MIME type string
        
        Args:
            mime_type: The MIME type string (e.g., 'text/plain', 'application/pdf')
        
        Returns:
            BlobType enum value or None if unknown MIME type
        """
        if not mime_type:
            return None
            
        mime_type = mime_type.lower()
        
        # Try to find a direct match with enum values
        for blob_type in cls:
            if blob_type.value.lower() == mime_type:
                return blob_type
                
        # Handle special cases and aliases
        mime_map = {
            # Text types with charset parameters
            'text/plain; charset=utf-8': cls.TEXT_PLAIN,
            'text/csv; charset=utf-8': cls.TEXT_CSV,
            'text/html; charset=utf-8': cls.TEXT_HTML,
            
            # Common aliases
            'application/javascript': cls.TEXT_JAVASCRIPT,
            'application/xhtml+xml': cls.TEXT_HTML,
            'text/xml; charset=utf-8': cls.TEXT_XML,
            'application/text': cls.TEXT_PLAIN,
            'text': cls.TEXT_PLAIN,
            
            # Office document types (both with and without parameters)
            'application/msword': cls.MS_WORD,
            'application/vnd.ms-word': cls.MS_WORD,
            'application/vnd.ms-excel': cls.MS_EXCEL,
            'application/excel': cls.MS_EXCEL,
            'application/vnd.ms-powerpoint': cls.MS_POWERPOINT,
            'application/powerpoint': cls.MS_POWERPOINT,
            
            # Image types
            'image/jpg': cls.IMAGE_JPEG,  # Common misspelling
            
            # Document types
            'application/x-pdf': cls.APP_PDF,
            
            # Archives
            'application/x-zip-compressed': cls.APP_ZIP,
            'application/x-gzip': cls.APP_GZIP
        }
        
        # Check the map for exact matches
        if mime_type in mime_map:
            return mime_map[mime_type]
            
        # Try to match just the main part before any parameters
        base_mime = mime_type.split(';')[0].strip()
        for blob_type in cls:
            if blob_type.value.lower() == base_mime:
                return blob_type
                
        # No match found
        return None

class Container:
    container_client: ContainerClient
    storage_account: StorageAccount
    # Reference to the BlobType enum for easier access
    BlobType: ClassVar[Enum] = BlobType

    def __init__(self, storage_account: StorageAccount, container_client: ContainerClient):
        self.storage_account = storage_account
        self.container_client = container_client

    def get_blob_names(self) -> List[str]:
        blobs = self.container_client.list_blob_names()
        return [blob for blob in blobs]
    
    def get_blobs(self) -> List[BlobProperties]:
        blobs = self.container_client.list_blobs()
        return [blob for blob in blobs]
        
    def get_blob_content(self, blob_name: str) -> Any:
        """
        Get the content of a specific blob using its name
        
        Args:
            blob_name: Name of the blob to retrieve
            
        Returns:
            bytes: The content of the blob
            
        Raises:
            ValueError: If the blob cannot be found or accessed
        """
        from io import BytesIO
        try:
            # Get the blob client for the specific blob
            blob_client = self.container_client.get_blob_client(blob_name)
            
            # Download the blob content
            download_stream = blob_client.download_blob()
            content = download_stream.readall()
        except Exception as e:
            raise ValueError(f"Error retrieving content for {blob_name = }: {str(e)}")
        
        blob_type = self.get_blob_type(blob_name)
        
        if blob_type is None:
            # Default to binary content if type cannot be determined
            return content
            
        # Process based on MIME type
        try:
            if blob_type in [BlobType.TEXT_PLAIN, BlobType.TEXT_CSV, BlobType.TEXT_HTML, 
                            BlobType.TEXT_CSS, BlobType.TEXT_JAVASCRIPT, BlobType.TEXT_XML, 
                            BlobType.TEXT_MARKDOWN, BlobType.APP_JSON, BlobType.APP_XML]:
                # Text-based formats
                text = content.decode('utf-8')
                return text
                
            elif blob_type == BlobType.MS_WORD:
                # Word documents
                try:
                    import docx                    
                    docx_file = BytesIO(content)
                    
                    # Load the document
                    document = docx.Document(docx_file)
                    
                    return document
                except ImportError:
                    raise ImportError("The python-docx package is required to read DOCX files. Install it with 'pip install python-docx'")
                except Exception as e:
                    raise ValueError(f"Error extracting text from DOCX blob {blob_name}: {str(e)}")
                
            elif blob_type == BlobType.APP_PDF:
                # PDF documents
                try:
                    import PyPDF2
                    pdf_file = BytesIO(content)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    
                    text = ""
                    for page_num in range(len(pdf_reader.pages)):
                        text += pdf_reader.pages[page_num].extract_text()
                    
                    return text
                except ImportError:
                    raise ImportError("The PyPDF2 package is required to read PDF files. Install it with 'pip install PyPDF2'")
                    
            elif blob_type == BlobType.MS_EXCEL:
                # Excel documents
                try:
                    import pandas as pd
                    excel_file = BytesIO(content)
                    return pd.read_excel(excel_file)
                except ImportError:
                    raise ImportError("The pandas package is required to read Excel files. Install it with 'pip install pandas openpyxl'")
                    
            else:
                # Binary content for all other types
                return content
                
        except Exception as e:
            raise ValueError(f"Error processing blob {blob_name} as {blob_type.name}: {str(e)}")
    
    def get_blob_content_by_properties(self, blob_properties: BlobProperties) -> bytes:
        """
        Get the content of a specific blob using its BlobProperties
        
        Args:
            blob_properties: BlobProperties object for the blob to retrieve
            
        Returns:
            bytes: The content of the blob
            
        Raises:
            ValueError: If the blob cannot be found or accessed
        """
        return self.get_blob_content(blob_properties.name)
    
    def get_blob_metadata(self) -> Dict[str, Dict[str, Any]]:
        """
        Retrieves metadata for all blobs in the container.
        Does NOT download content or generate hash.

        Returns:
            Dict mapping blob name to its metadata (last_modified, size, etag).
        """
        print(f"Retrieving metadata for blobs in container '{self.container_client.container_name}'...")
        blob_metadata = {}
        blobs: List[BlobProperties] = self.get_blobs() # Use self.get_blobs()
        print(f"Found {len(blobs)} blobs in container.")
        for blob_prop in blobs:
            # Ensure last_modified is timezone-aware (UTC)
            last_modified_utc = blob_prop.last_modified.astimezone(timezone.utc)

            blob_metadata[blob_prop.name] = {
                'name': blob_prop.name,
                'last_modified': last_modified_utc,
                'size': blob_prop.size,
                'etag': blob_prop.etag
            }
        return blob_metadata
        
    def get_docx_content(self, blob_name: str) -> str:
        """
        Get the content of a Word DOCX file as text
        
        Args:
            blob_name: Name of the blob to retrieve
            
        Returns:
            str: The text content extracted from the DOCX file
            
        Raises:
            ValueError: If the blob cannot be found, accessed, or is not a valid DOCX file
            ImportError: If the required docx package is not installed
        """
        try:
            # Import docx library for processing Word documents
            try:
                import docx
            except ImportError:
                raise ImportError("The python-docx package is required to read DOCX files. Install it with 'pip install python-docx'")
            
            # Get the blob content as bytes
            content_bytes = self.get_blob_content(blob_name)
            
            # Create a file-like object from the bytes
            from io import BytesIO
            docx_file = BytesIO(content_bytes)
            
            # Load the document
            document = docx.Document(docx_file)
            
            # Extract text from all paragraphs
            paragraphs = [para.text for para in document.paragraphs]
            
            # Join paragraphs with newlines
            text_content = '\n'.join(paragraphs)
            
            return text_content
        except ImportError as e:
            raise e
        except Exception as e:
            print (f"Error extracting text from DOCX blob {blob_name}: {str(e)}")
            raise e
        
    def delete_blob(self, blob_name: str) -> bool:
        """
        Delete a blob from the container by its name
        
        Args:
            blob_name: Name of the blob to delete
            
        Returns:
            bool: True if the blob was successfully deleted, False otherwise
            
        Raises:
            ValueError: If there's an error deleting the blob
        """
        try:
            # Get the blob client for the specific blob
            blob_client = self.container_client.get_blob_client(blob_name)
            
            # Delete the blob
            blob_client.delete_blob()
            
            return True
        except Exception as e:
            print (f"Error deleting blob {blob_name}: {str(e)}")
            raise e
            
    def find_blobs_by_filename(self, filename: str, case_sensitive: bool = False) -> List[BlobProperties]:
        """
        Find all blobs that match a particular filename, regardless of path
        
        Args:
            filename: The filename to search for (without path)
            case_sensitive: Whether the search should be case-sensitive
            
        Returns:
            List[BlobProperties]: List of BlobProperties objects for matching blobs
            
        Examples:
            # Find all files named 'data.csv' in any folder
            data_files = container.find_blobs_by_filename('data.csv')
            
            # Find all JSON files
            json_files = container.find_blobs_by_filename('.json', case_sensitive=False)
        """
        try:
            # Get all blobs in the container
            all_blobs = self.container_client.list_blobs()
            
            # Filter blobs by filename
            matching_blobs = []
            
            for blob in all_blobs:
                # Extract just the filename part (after the last '/')
                blob_full_name = blob.name
                blob_filename = blob_full_name.split('/')[-1]
                
                # Check if the filename matches
                if case_sensitive:
                    if filename in blob_filename:
                        matching_blobs.append(blob)
                else:
                    if filename.lower() in blob_filename.lower():
                        matching_blobs.append(blob)
            
            return matching_blobs
        except Exception as e:
            print(f"Error searching for blobs with filename '{filename}': {str(e)}")
            raise e
                       
    def get_blob_type_from_properties(self, properties: BlobProperties) -> Optional["BlobType"]:
        """
        Detect the MIME type of a blob based on its properties
        
        Args:
            properties: BlobProperties object to analyze
            
        Returns:
            BlobType: The detected MIME type or None if cannot be determined
        """
        # First try from content type
        if properties and properties.content_settings and properties.content_settings.content_type:
            content_type = properties.content_settings.content_type
            blob_type = BlobType.from_mime_type(content_type)
            if blob_type:
                return blob_type
                
        # Fall back to extension-based detection using the blob name
        if properties and properties.name and '.' in properties.name:
            extension = properties.name.split('.')[-1]
            return BlobType.from_extension(extension)
            
        return None
            
    def get_blob_type(self, blob_name: str) -> Optional["BlobType"]:
        """
        Detect the MIME type of a blob based on its content type and/or extension
        
        Args:
            blob_name: Name of the blob to analyze
            
        Returns:
            BlobType: The detected MIME type or None if cannot be determined
        """
        try:
            # First try to get the content type from blob properties
            blob_client = self.container_client.get_blob_client(blob_name)
            properties = blob_client.get_blob_properties()
            return self.get_blob_type_from_properties(properties)
        except Exception:
            # If we can't get properties, fall back to extension-based detection
            pass
            
        # Fall back to extension-based detection
        if '.' not in blob_name:
            return None
            
        extension = blob_name.split('.')[-1]
        return BlobType.from_extension(extension)
        
    def process_blob_by_type(self, blob_name: str) -> Any:
        """
        Process a blob based on its detected type
        
        Args:
            blob_name: Name of the blob to process
            
        Returns:
            The processed content in the appropriate format for the detected type
            
        Raises:
            ValueError: If the blob cannot be processed
            ImportError: If a required package is not installed
        """
        blob_type = self.get_blob_type(blob_name)
        
        if blob_type is None:
            # Default to binary content if type cannot be determined
            return self.get_blob_content(blob_name)
            
        # Process based on MIME type
        try:
            if blob_type in [BlobType.TEXT_PLAIN, BlobType.TEXT_CSV, BlobType.TEXT_HTML, 
                            BlobType.TEXT_CSS, BlobType.TEXT_JAVASCRIPT, BlobType.TEXT_XML, 
                            BlobType.TEXT_MARKDOWN, BlobType.APP_JSON, BlobType.APP_XML]:
                # Text-based formats
                return self.get_text_content(blob_name)
                
            elif blob_type == BlobType.MS_WORD:
                # Word documents
                return self.get_docx_content(blob_name)
                
            elif blob_type == BlobType.APP_PDF:
                # PDF documents
                try:
                    import PyPDF2
                    content_bytes = self.get_blob_content(blob_name)
                    from io import BytesIO
                    pdf_file = BytesIO(content_bytes)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    
                    text = ""
                    for page_num in range(len(pdf_reader.pages)):
                        text += pdf_reader.pages[page_num].extract_text()
                    
                    return text
                except ImportError:
                    raise ImportError("The PyPDF2 package is required to read PDF files. Install it with 'pip install PyPDF2'")
                    
            elif blob_type == BlobType.MS_EXCEL:
                # Excel documents
                try:
                    import pandas as pd
                    content_bytes = self.get_blob_content(blob_name)
                    from io import BytesIO
                    excel_file = BytesIO(content_bytes)
                    return pd.read_excel(excel_file)
                except ImportError:
                    raise ImportError("The pandas package is required to read Excel files. Install it with 'pip install pandas openpyxl'")
                    
            else:
                # Binary content for all other types
                return self.get_blob_content(blob_name)
                
        except Exception as e:
            print(f"Error processing blob {blob_name} as {blob_type.name}: {str(e)}")
            raise e

    def get_folder_structure(self) -> Dict[str, List[str]]:
        """
        Get the folder structure and files in a container
        
        Args:
            container: The blob container
            
        Returns:
            Dict[str, List[str]]: Dictionary with folders as keys and lists of files as values
        """
        blobs: List[BlobProperties] = self.container_client.list_blobs()
        
        # Dictionary to store folder structure
        folder_structure = {}
        
        for blob in blobs:
            # Get the blob name
            blob_name = blob.name
            
            # Determine the folder
            if '/' in blob_name:
                folder_path = '/'.join(blob_name.split('/')[:-1])
                file_name = blob_name.split('/')[-1]
                
                # Add folder to dictionary if it doesn't exist
                if folder_path not in folder_structure:
                    folder_structure[folder_path] = []
                
                # Add file to folder
                folder_structure[folder_path].append(file_name)
            else:
                # Files in the root directory
                if 'root' not in folder_structure:
                    folder_structure['root'] = []
                
                folder_structure['root'].append(blob_name)
        
        return folder_structure
    
    def resolve_content_type( file_path:str, content_type: Optional[Union[str, "BlobType"]] = None) -> Optional[str]:
        resolved_content_type = None
        if content_type:
            if isinstance(content_type, BlobType):
                resolved_content_type = content_type.value
            else:
                resolved_content_type = content_type
        else: 
            # Try to determine from file extension
            _, extension = os.path.splitext(file_path)
            if extension:
                blob_type = BlobType.from_extension(extension)
                if blob_type:
                    resolved_content_type = blob_type.value
        return resolved_content_type
    
    from io import BufferedReader
    def upload_stream(self, file_data:BufferedReader, destination_blob_name: Optional[str] = None, 
                    resolved_content_type: str = None, 
                    metadata: Optional[Dict[str, str]] = None) -> BlobProperties:

        # Get the blob client for the destination
        blob_client = self.container_client.get_blob_client(destination_blob_name)
        
        # Set up content settings
        content_settings = None
        if resolved_content_type:
            from azure.storage.blob import ContentSettings
            content_settings = ContentSettings(content_type=resolved_content_type)
        
        # Read the file and upload
        blob_client.upload_blob(
            file_data, 
            overwrite=True,
            content_settings=content_settings,
            metadata=metadata
        )
        
        # Return properties of the uploaded blob
        return blob_client.get_blob_properties()

    def upload_file(self, local_file_path: str, destination_blob_name: Optional[str] = None, 
                    content_type: Optional[Union[str, "BlobType"]] = None, 
                    metadata: Optional[Dict[str, str]] = None) -> BlobProperties:
        """
        Upload a file from a local path to the container.
        
        Args:
            local_file_path: Path to the local file to upload
            destination_blob_name: Name to use for the blob in the container (defaults to filename)
            content_type: Content type to set for the blob (can be string or BlobType enum)
            metadata: Optional dictionary of custom metadata to set on the blob
        
        Returns:
            BlobProperties: Properties of the uploaded blob
            
        Raises:
            ValueError: If the file does not exist or cannot be uploaded
            
        Examples:
            # Upload a file with default name (same as local filename)
            container.upload_file("C:/documents/report.pdf")
            
            # Upload a file with a custom name and content type
            container.upload_file("C:/images/photo.jpg", "vacation/beach.jpg", BlobType.IMAGE_JPEG)
            
            # Upload with metadata
            container.upload_file("C:/data/records.csv", "2023/monthly_data.csv", 
                                BlobType.TEXT_CSV, {"source": "finance", "period": "Q2"})
        """
        # Verify the file exists
        if not os.path.isfile(local_file_path):
            raise ValueError(f"File does not exist: {local_file_path}")
        
        # If destination blob name not provided, use the filename
        if destination_blob_name is None:
            destination_blob_name = os.path.basename(local_file_path)
        
        # Determine content type if not explicitly provided
        resolved_content_type = Container.resolve_content_type(local_file_path, content_type)
        
        # Set up content settings
        content_settings = None
        if resolved_content_type:
            from azure.storage.blob import ContentSettings
            content_settings = ContentSettings(content_type=resolved_content_type)
        
        # Read the file and upload
        with open(local_file_path, "rb") as file_data:
            blob_properties = self.upload_stream( file_data=file_data, destination_blob_name=destination_blob_name,
                               resolved_content_type=resolved_content_type,
                               metadata=metadata)

        # Return properties of the uploaded blob
        return blob_properties

class FolderProcessingResults:
    file_mappings: Dict[str, str]
    uuid_to_original: Dict[str, str]
    def __init__(self, file_mappings: Dict[str, str], uuid_to_original: Dict[str, str]):
        self.file_mappings = file_mappings
        self.uuid_to_original = uuid_to_original

    import json
    def store_mapping(self, mapping_file: str) -> None:
        with open(mapping_file, 'w') as f:
            json.dump(self.file_mappings, f, indent=2)
    
    def store_uuid_mapping(self, uuid_mapping_file: str) -> None:
        with open(uuid_mapping_file, 'w') as f:
            json.dump(self.uuid_to_original, f, indent=2)
    
    def load_file_mapping(self, mapping_file: str) -> None:
        with open(mapping_file, 'r') as f:
            mappings = json.load(f)
        self.file_mappings = mappings
    
    def load_uuid_mapping(self, uuid_mapping_file: str) -> None:
        with open(uuid_mapping_file, 'r') as f:
            uuid_mappings = json.load(f)
        self.uuid_to_original = uuid_mappings

    def load_mappings(self, mapping_file: str, uuid_mapping_file: str) -> None:
        self.load_file_mapping(mapping_file)
        self.load_uuid_mapping(uuid_mapping_file)

    def get_uuid_filename(self, rel_filename: str) -> Optional[str]:
        return self.file_mappings.get(rel_filename, None)
    
    def get_original_filename(self, uuid_filename: str) -> Optional[str]:
        return self.uuid_to_original.get(uuid_filename, None)

class FolderProcessor:
    localfolder: str

    def __init__(self, folder: str):
        # Validate the local directory
        if not os.path.isdir(folder):
            raise ValueError(f"The directory '{folder}' does not exist or is not a directory")
        self.localfolder = folder    

    def get_uuid_filename( filename: str) -> str:
        import uuid
        # Get file extension
        file_name, file_ext = os.path.splitext(filename)
        if file_name == str(uuid.UUID(file_name)):
            # If the filename is already a UUID, keep it as is
            uuid_filename = file_name + file_ext
        else:
            # Generate a UUID and keep the file extension
            uuid_filename = f"{uuid.uuid4()}{file_ext}"
        return uuid_filename

    def get_dest_path(container_folder:str, rel_path: str, uuid_filename:str) -> str :
        # Keep folder structure: container_folder + relative_path (with UUID filename)
        rel_dir = os.path.dirname(rel_path)
        if rel_dir:
            # Replace backslashes with forward slashes for blob paths
            rel_dir = rel_dir.replace('\\', '/')
            dest_path = f"{container_folder}{rel_dir}/{uuid_filename}"
        else:
            dest_path = f"{container_folder}{uuid_filename}"    
        return dest_path

    #from typing import Dict, List, Optional, Tuple

    def upload_directory(self, 
                        container: Container,
                        container_folder: str = "", 
                        maintain_structure: bool = False, 
                        extensions_filter: Optional[List[str]] = None,
                        use_uuid: bool = True,
                        overwrite: bool = False,
                        record_mappings: bool = True) -> Dict[str, str]:
        """
        Upload all files from a local directory to the container, either preserving the folder 
        structure or flattening all files into a single container folder.
        
        Args:
            local_directory: Path to the local directory to upload
            container_folder: Base folder in the container to upload to (will be created if doesn't exist)
            maintain_structure: If True, preserves folder structure; if False, flattens all files
            extensions_filter: Optional list of file extensions to include (e.g., ['.pdf', '.docx'])
            use_uuid: If True, replaces filenames with UUIDs; if False, keeps original filenames
            overwrite: Whether to overwrite if files already exist
            record_mappings: Whether to create a mapping file relating original names to UUIDs
            
        Returns:
            Dict mapping original file paths to their destination blob paths
        
        Raises:
            ValueError: If the local directory doesn't exist or if upload fails
        """
        
        # Normalize the container folder (ensure it ends with a slash if not empty)
        if container_folder and not container_folder.endswith('/'):
            container_folder = container_folder + '/'
        
        # Initialize mapping of original files to blob paths
        file_mappings = {}
        
        # Keep a mapping of UUIDs to original filenames if using UUIDs
        uuid_to_original = {}

                
        # Walk through the directory structure
        for root, _, files in os.walk(self.local_directory):
            # Filter files by extension if specified
            if extensions_filter:
                files = [f for f in files if any(f.lower().endswith(ext.lower()) for ext in extensions_filter)]
            
            for filename in files:
                # Get the full local file path
                local_file_path = os.path.join(root, filename)
                # Determine relative path from the base directory
                rel_path = os.path.relpath(local_file_path, self.local_directory)
                # Generate a UUID for the file if requested
                if use_uuid:
                    # Store mapping
                    uuid_to_original[FolderProcessor.get_uuid_filename(filename)] = rel_path
                else:
                    uuid_filename = filename
                
                # Determine the destination blob path
                if maintain_structure:
                    # Keep folder structure: container_folder + relative_path (with UUID filename)
                    rel_dir = os.path.dirname(rel_path)
                    dest_path = FolderProcessor.get_dest_path(container_folder, rel_path, uuid_filename)
                else:
                    # Flatten structure: all files go directly to container_folder
                    dest_path = f"{container_folder}{uuid_filename}"
                
                import mimetypes
                # Detect content type
                content_type, _ = mimetypes.guess_type(local_file_path)
                
                # Create a blob client for the destination path
                blob_client = self.container_client.get_blob_client(dest_path)
                
                # Set content settings if content type was detected
                content_settings = None
                if content_type:
                    from azure.storage.blob import ContentSettings
                    content_settings = ContentSettings(content_type=content_type)
                
                # Upload the file
                with open(local_file_path, "rb") as data:
                    blob_client.upload_blob(data, overwrite=overwrite, content_settings=content_settings)
                
                # Store the mapping
                file_mappings[rel_path] = dest_path
                
                print(f"Uploaded: {rel_path} -> {dest_path}")
        
        # If using UUIDs, create a mapping file in the container
        if use_uuid and record_mappings:
            # Create a JSON mapping file
            mapping_blob_name = f"{container_folder}file_mappings.json"
            mapping_blob = self.container_client.get_blob_client(mapping_blob_name)
            
            # Convert mapping to JSON
            mapping_json = json.dumps(uuid_to_original, indent=2)
            
            # Upload the mapping file
            mapping_blob.upload_blob(mapping_json, overwrite=True, content_settings=ContentSettings(content_type="application/json"))
            print(f"Created mapping file: {mapping_blob_name}")
        
        return FolderProcessingResults( file_mappings=file_mappings, uuid_to_original=uuid_to_original)

    # def download_files_with_mapping(self, 
    #                             container_folder: str = "",
    #                             mapping_file: Optional[str] = None,
    #                             use_original_names: bool = True) -> List[str]:
    #     """
    #     Download files from the container, optionally using a mapping file to restore original names.
        
    #     Args:
    #         container_folder: Folder in the container to download from
    #         mapping_file: Optional path to the mapping file (defaults to 'file_mappings.json' in container_folder)
    #         use_original_names: If True, restores original filenames using the mapping file
            
    #     Returns:
    #         List of local file paths that were downloaded
            
    #     Raises:
    #         ValueError: If download fails or mapping file not found when use_original_names=True
    #     """
    #     # Ensure the local directory exists
    #     os.makedirs(self.local_directory, exist_ok=True)
        
    #     # Normalize the container folder
    #     if container_folder and not container_folder.endswith('/'):
    #         container_folder = container_folder + '/'
        
    #     # Initialize list of downloaded files
    #     downloaded_files = []
        
    #     # If we need to use original names, load the mapping file
    #     uuid_to_original_map = {}
    #     if use_original_names:
    #         # Determine mapping file path if not specified
    #         if not mapping_file:
    #             mapping_file = f"{container_folder}file_mappings.json"
            
    #         # Get the mapping file
    #         mapping_blob = self.container_client.get_blob_client(mapping_file)
    #         try:
    #             mapping_content = mapping_blob.download_blob().readall()
    #             uuid_to_original_map = json.loads(mapping_content)
    #         except Exception as e:
    #             raise ValueError(f"Failed to load mapping file: {str(e)}")
        
    #     # List all blobs in the specified folder
    #     blobs = self.container_client.list_blobs(name_starts_with=container_folder)
        
    #     for blob in blobs:
    #         # Skip the mapping file itself
    #         if blob.name == mapping_file:
    #             continue
                
    #         # Get the blob name (without container folder prefix)
    #         if container_folder:
    #             relative_blob_name = blob.name[len(container_folder):]
    #         else:
    #             relative_blob_name = blob.name
                
    #         # Skip if empty (this would be the container folder itself)
    #         if not relative_blob_name:
    #             continue
                
    #         # Determine the local file path
    #         if use_original_names:
    #             # Check if this blob has a mapping entry
    #             blob_filename = os.path.basename(blob.name)
    #             if blob_filename in uuid_to_original_map:
    #                 # Use the original path from the mapping
    #                 original_path = uuid_to_original_map[blob_filename]
    #                 local_file_path = os.path.join(self.local_directory, original_path)
    #             else:
    #                 # If no mapping found, use the blob name as is
    #                 local_file_path = os.path.join(self.local_directory, relative_blob_name)
    #         else:
    #             # Use the blob path directly
    #             local_file_path = os.path.join(self.local_directory, relative_blob_name)
            
    #         # Ensure the directory exists
    #         os.makedirs(os.path.dirname(local_file_path), exist_ok=True)
            
    #         # Download the blob
    #         blob_client = self.container_client.get_blob_client(blob.name)
    #         with open(local_file_path, "wb") as file:
    #             download_stream = blob_client.download_blob()
    #             file.write(download_stream.readall())
            
    #         downloaded_files.append(local_file_path)
    #         print(f"Downloaded: {blob.name} -> {local_file_path}")
        
    #     return downloaded_files

import azure.search.documents.indexes as azsdi
import azure.search.documents.indexes.models as azsdim
from azure.search.documents.indexes import SearchIndexerClient
from azure.core.credentials import AzureKeyCredential
from docx import Document
from io import BytesIO

from openai import AzureOpenAI
from langchain_openai import AzureChatOpenAI
class AIService:
    cognitive_client: CognitiveServicesManagementClient
    azure_account: azcsm.Account
    resource_group: ResourceGroup
    
    
    def __init__(self, 
                 resource_group: ResourceGroup,
                 cognitive_client: CognitiveServicesManagementClient,
                 azure_Account: azcsm.Account):
        self.resource_group = resource_group
        self.cognitive_client = cognitive_client
        self.azure_account = azure_Account
        
    def get_OpenAIClient(self, api_version:str) -> "OpenAIClient" :
        keys = self.cognitive_client.accounts.list_keys(self.resource_group.get_name(), self.azure_account.name)
        openai_client = AzureOpenAI(
            api_key=keys.key1,
            api_version=api_version,
            azure_endpoint= f"https://{self.azure_account.name}.openai.azure.com/",
        )
        return OpenAIClient(self, openai_client) 

    def get_models(self, azure_location: str = None) -> List[azcsm.Model]:
        if (azure_location is None): 
            azure_location = self.resource_group.azure_resource_group.location
        models_list = self.cognitive_client.models.list(azure_location)
        models = [model for model in models_list]
        return models 

    @staticmethod
    def get_model_details(model: azcsm.Model) -> Dict[str, Any]:
        """
        Get details for a specific model.
        
        Args:
            model_id: The model to be processed
            
        Returns:
            Dictionary with model details
        """
        try:
            info =  {
                "kind": model.kind,
                "name": model.model.name,
                "format": model.model.format,
                "version": model.model.version,
                "sju_name": model.sku_name,
            }
            return info
        except Exception as e:
            print(f"Error getting model '{model.model.name}': {str(e)}")
            return {}

    def get_deployments(self) -> List[azcsm.Deployment]:
        try:
            deployments = list(self.cognitive_client.deployments.list(self.resource_group.get_name(), self.azure_account.name))
            
            result = [ deployment for deployment in deployments ]
            return result
        except Exception as e:
            print(f"Error listing deployments: {str(e)}")
            return []
        
    def get_deployment(self, deployment_name:str) -> azcsm.Deployment : 
            deployment = self.cognitive_client.deployments.get( resource_group_name=self.resource_group.get_name(), account_name=self.azure_account.name, deployment_name=deployment_name )        
            return deployment

    @staticmethod        
    def get_deployment_details(deployment: azcsm.Deployment) -> Dict[str, Any]:
        """
        Get details for a specific deployment.
        
        Args:
            deployment: The deployment
            
        Returns:
            Dictionary with deployment details
        """
        try:
            # Handle missing properties safely
            deployment_info = {
                "name": deployment.name if hasattr(deployment, 'name') else "name not found",
                "status": "unknown"
            }
            # Add properties only if they exist
            if hasattr(deployment, 'properties'):
                props = deployment.properties
                if hasattr(props, 'model'):
                    deployment_info["model"] = props.model
                if hasattr(props, 'provisioning_state'):
                    deployment_info["status"] = props.provisioning_state
                # Handle scale settings if they exist
                if hasattr(props, 'scale_settings') and props.scale_settings is not None:
                    scale_settings = {}
                    if hasattr(props.scale_settings, 'scale_type'):
                        scale_settings["scale_type"] = props.scale_settings.scale_type
                    if hasattr(props.scale_settings, 'capacity'):
                        scale_settings["capacity"] = props.scale_settings.capacity
                    deployment_info["scale_settings"] = scale_settings
                # Add timestamps if they exist
                if hasattr(props, 'created_at'):
                    deployment_info["created_at"] = props.created_at
                if hasattr(props, 'last_modified'):
                    deployment_info["last_modified"] = props.last_modified
            return deployment_info
            
        except Exception as e:
            print(f"Error getting deployment '{deployment.name}': {str(e)}")
            return {}

    def create_deployment(self, 
                         deployment_name: str, 
                         model_name: str, 
                         model_version:str = None,
                         sku_name: str = "Standard",
                         capacity: int = 1) -> Union[azcsm.Deployment, Dict[str, str]]:
        """
        Create a new model deployment in Azure OpenAI.
        
        Args:
            deployment_name: Name for the new deployment
            model: Base model name (e.g., 'gpt-4', 'text-embedding-ada-002')
            capacity: Number of tokens per minute in millions (TPM)
            scale_type: Scaling type (Standard, Manual)
            
        Returns:
            the Deployment when prepared
        """

        try:
            if model_version:
                model = azcsm.Model(name=model_name, version=model_version)
            else:
                model = azcsm.Model(name=model_name)

            deployment_properties = azcsm.DeploymentProperties(model=model)
            
            # Create SKU configuration
            sku = azcsm.Sku(name=sku_name, capacity=capacity)

            # properties = azcsm.DeploymentProperties(

            #     model=model,
            #     scale_settings=azcsm.DeploymentScaleSettings(
            #         scale_type=scale_type,
            #         capacity=capacity
            #     ),
            #     rai_policy_name="Microsoft.Default"
            # )
            poller = self.cognitive_client.deployments.begin_create_or_update(
                resource_group_name=self.resource_group.get_name(),
                account_name=self.azure_account.name,
                deployment_name=deployment_name,
                deployment=None, 
                parameters = { 
                    "properties": deployment_properties,
                    "sku": sku
                }
            )
            deployment: azcsm.Deployment = poller.result()
            return deployment
            
        except Exception as e:
            print(f"Error creating deployment '{deployment_name}': {str(e)}")
            return {"error": str(e)}

    def delete_deployment(self, deployment_name: str) -> bool:
        """
        Delete a model deployment in Azure OpenAI.
        
        Args:
            deployment_name: Name of the deployment to delete
            
        Returns:
            Boolean indicating success or failure
        """
        try:
            # Start the delete operation
            poller = self.cognitive_client.deployments.begin_delete(
                resource_group_name=self.resource_group.get_name(),
                account_name=self.azure_account.name,
                deployment_name=deployment_name
            )
            
            # Wait for the operation to complete
            result = poller.result()

            print(f"Successfully deleted deployment '{deployment_name}'")
            return True
        except Exception as e:
            print(f"Error deleting deployment '{deployment_name}': {str(e)}")
            return False

    def update_deployment(self, 
                         deployment_name: str, 
                         sku_name: str = "Standard",
                         capacity: int = 1) -> Union[azcsm.Deployment, Dict[str, str]]:
        """
        Update an existing model deployment in Azure OpenAI.
        
        Args:
            deployment_name: Name of the deployment to update
            capacity: New capacity value (optional)
            scale_type: New scale type (optional)
            
        Returns:
            Dictionary with deployment details
        """
        try:
            deployment = self.get_deployment(deployment_name)

            model_props = deployment.properties.model
            model = azcsm.Model(
                # If the model is stored as a complex object, preserve its name/version
                name=model_props.name if hasattr(model_props, 'name') else model_props,
                version=model_props.version if hasattr(model_props, 'version') else None
            )            
            updated_sku = azcsm.Sku(name=sku_name, capacity=capacity)

            deployment_properties = azcsm.DeploymentProperties(model=model)

            poller = self.cognitive_client.deployments.begin_create_or_update(
                self.resource_group.get_name(),
                self.azure_account.name,
                deployment_name,
                parameters={
                    "properties": deployment_properties,
                    "sku": updated_sku
                }
            )            
            deployment = poller.result()
            return deployment
        except Exception as e:
            print(f"Error updating deployment '{deployment_name}': {str(e)}")
            return {"error": str(e)}

class OpenAIClient:
    ai_service: AIService 
    openai_client: AzureOpenAI

    def __init__(self, ai_service: AIService, openai_client: AzureOpenAI):
        self.ai_service = ai_service
        self.openai_client = openai_client
    
    @retry(wait=wait_random_exponential(min=1, max=20), stop=stop_after_attempt(6))
    def generate_embeddings(self, text: str, model: str = "text-embedding-3-large") -> List[float]:
        """
        Generate embeddings for text using Azure OpenAI.
        
        Args:
            text: The text to generate embeddings for
            model: The embedding model to use
            
        Returns:
            List of float values representing the embedding vector
        """
        try:
            response = self.openai_client.embeddings.create(input=text, model=model)
            return response.data[0].embedding
        except Exception as e:
            print(f"Error generating embeddings: {str(e)}")
            raise e
    
    def generate_chat_completion(self, 
                                messages: List[Dict[str, str]], 
                                model: str, 
                                temperature: float = 0.7, 
                                max_tokens: int = 800,
                                response_format: Dict[str, Any] = None,
                                ) -> Dict[str, Any]:
        """
        Generate a chat completion using Azure OpenAI.
        
        Args:
            messages: List of message dictionaries with 'role' and 'content'
            model: The deployment name of the model
            temperature: Temperature for generation (0-1)
            max_tokens: Maximum number of tokens to generate
            
        Returns:
            Chat completion response
        """
        try:
            response = self.openai_client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
                response_format=response_format
            )
            return {
                "content": response.choices[0].message.content,
                "finish_reason": response.choices[0].finish_reason,
                "usage": {
                    "prompt_tokens": response.usage.prompt_tokens,
                    "completion_tokens": response.usage.completion_tokens,
                    "total_tokens": response.usage.total_tokens
                }
            }
        except Exception as e:
            print(f"Error generating chat completion: {str(e)}")
            return {"error": str(e)}

class SearchService:
    search_service: azsrm.SearchService
    subscription: Subscription
    resource_group: ResourceGroup
    index_client: azsdi.SearchIndexClient
    search_client: azsd.SearchClient
    openai_client: Any
    index_name: str

    def __init__(self, subscription: Subscription, resource_group: ResourceGroup, search_service: azsrm.SearchService):
        self.subscription = subscription
        self.resource_group = resource_group
        self.search_service = search_service
        self.index_client: Optional[azsdi.SearchIndexClient] = None
        self.search_client: Optional[azsd.SearchClient] = None
        self.openai_client: Optional[Any] = None
        self.index_name: str = os.getenv("INDEX_NAME", "default-index")
    
    def get_admin_key(self) -> str:
        search_mgmt_client = SearchManagementClient(self.resource_group.subscription.identity.get_credential(),
                                                    self.resource_group.subscription.subscription_id)
        keys = search_mgmt_client.admin_keys.get(resource_group_name=self.resource_group.azure_resource_group.name,
                                                search_service_name=self.search_service.name)
        return keys.primary_key

    def get_credential(self) -> AzureKeyCredential:
        return AzureKeyCredential(self.get_admin_key())

    def get_service_endpoint(self) -> str:
        return f"https://{self.search_service.name}.search.windows.net"

    def get_index_client(self) -> azsdi.SearchIndexClient:
        if self.index_client is None:
            self.index_client = azsdi.SearchIndexClient(self.get_service_endpoint(),
                                                         self.get_credential())
        return self.index_client
    
    def get_indexes(self) -> List[azsdim.SearchIndex]:
        index_client = self.get_index_client()
        indexes = list(index_client.list_indexes())
        return indexes

    def get_index(self, index_name: str) -> Optional["SearchIndex"]:
        index_client = self.get_index_client()
        index = index_client.get_index(index_name)
        if index and index.name == index_name:
            return SearchIndex(self, index.name, index.fields, index.vector_search)
        return None
    
    def add_simple_field(self, field_name: str, field_type: str, filterable: bool = False, is_key: bool = False, **kw) -> azsdim.SimpleField:
        """
        Create a new field for the search index.
        
        Args:
            field_name: The name of the field
            field_type: The type of the field as defined in azure.search.documents.indexes.models.SearchFieldDataType
            is_filterable: Whether the field is filterable
            is_key: Whether the field uniquely identifies documents in the index
            
        Returns:
            SimpleField: The created simple field
        """

        from azure.search.documents.indexes.models import SearchFieldDataType

        resolved_type = getattr(SearchFieldDataType, field_type)

        return azsdim.SimpleField(name=field_name, type=resolved_type, filterable=filterable, key=is_key, **kw)
    
    def add_searchable_field(self, field_name: str, field_collection_type: bool = False, filterable: bool = False, searchable: bool = True, is_key: bool = False, analyzer_name: Optional[Union[str, str]] = None, **kw) -> azsdim.SearchableField:
        """
        Create a new searchable field for the search index.
        
        Args:
            field_name: The name of the field
            field_collection_type: Whether the field is a collection or a String
            is_filterable: Whether the field is filterable
            is_searchable: Whether the field is searchable
            is_key: Whether the field uniquely identifies documents in the index
            analyzer_name: The name of the analyzer to use for this field
            
        Returns:
            SearchableField: The created searchable field
        """
        return azsdim.SearchableField(name=field_name, collection=field_collection_type, searchable=searchable, filterable=filterable, key=is_key, analyzer_name=analyzer_name, **kw)
    
    def add_search_field(self, field_name: str, field_type: str, searchable: bool = True, vector_search_dimensions: Optional[int] = None, vector_search_profile_name: Optional[str] = None, **kw) -> azsdim.SearchField:
        """
        Create a new search field for the search index.
        
        Args:
            field_name: The name of the field
            field_type: The type of the field
            is_searchable: Whether the field is searchable
            vector_search_dimensions: The dimensions for the space used for vector search
            vector_search_profile_name: The name of the vector search profile
            
        Returns:
            SearchField: The created search field
        """
        return azsdim.SearchField(name=field_name, type=field_type, searchable=searchable, vector_search_dimensions=vector_search_dimensions, vector_search_profile_name=vector_search_profile_name, **kw)

    def create_or_update_index(self, index_name: str, fields: List[azsdim.SearchField], vector_search: Optional[azsdim.VectorSearch] = None) -> "SearchIndex":
        return SearchIndex(self, index_name, fields, vector_search)
    
    def add_semantic_configuration(self,
                                   index_name: str,
                                   title_field: str = "title",
                                   content_fields: Optional[List[str]] = None,
                                   keyword_fields: Optional[List[str]] = None,
                                   semantic_config_name: str = "default-semantic-config") -> azsdim.SearchIndex:
        """
        Add semantic configuration to the index.
        
        Args:
            index_name: The name of the index to be modified
            title_field: The name of the title field
            content_fields: List of content fields to prioritize
            keyword_fields: List of keyword fields to prioritize
            semantic_config_name: The name of the semantic configuration
            
        Returns:
            The updated index
        """
        if content_fields is None:
            content_fields = ["content"]
        
        if keyword_fields is None:
            keyword_fields = ["tags"]
        
        # Get the existing index
        index = self.get_index_client().get_index(index_name)
        
        # Define semantic configuration
        semantic_config = azsdim.SemanticConfiguration(
            name=semantic_config_name,
            prioritized_fields=azsdim.SemanticPrioritizedFields(
                title_field=azsdim.SemanticField(field_name=title_field),
                content_fields=[
                    azsdim.SemanticField(field_name=field) for field in content_fields
                ],
                keywords_fields=[
                    azsdim.SemanticField(field_name=field) for field in keyword_fields
                ]
            )
        )
        
        # Create SemanticSearch instance
        semantic_search = azsdim.SemanticSearch(configurations=[semantic_config])
        
        # Add semantic settings to the index
        index.semantic_search = semantic_search
        
        # Update the index
        result = self.get_index_client().create_or_update_index(index)
        
    def get_indexer_client(self) -> SearchIndexerClient:
        """
        Get a SearchIndexerClient for this search service.
        
        Returns:
            SearchIndexerClient: A client to interact with Azure search service Indexers
        """
        return SearchIndexerClient(
            endpoint=self.get_service_endpoint(),
            credential=self.get_credential()
        )
    
    def create_indexer_manager(self) -> "SearchIndexerManager":
        """
        Create a SearchIndexerManager for this search service.
        
        Returns:
            SearchIndexerManager: A manager for working with indexers, data sources, and skillsets
        """
        return SearchIndexerManager(self)
    
def get_std_vector_search( connections_per_node:int = 4, 
                          neighbors_list_size: int = 400, 
                          search_list_size: int = 500, 
                          metric: str = "cosine") -> azsdim.VectorSearch:
    """
    Get a standard vector search configuration.
    Args:
        connections_per_node: Number of connections per node. Default is 4.
        neighbors_list_size: Size of the dynamic list for nearest neighbors. Default is 400.
        search_list_size: Size of the dynamic list for searching. Default is 500.
        metric: Distance metric (cosine, euclidean, dotProduct). Default is cosine.
    Returns:
        A vector search configuration
    """
    # Define vector search configuration
    vector_search = azsdim.VectorSearch(
        algorithms=[
            azsdim.VectorSearchAlgorithmConfiguration(
                name="default-algorithm",
                kind="hnsw",
                hnsw_parameters=azsdim.HnswParameters(
                    m=4,  # Number of connections per node
                    ef_construction=400,  # Size of the dynamic list for nearest neighbors
                    ef_search=500,  # Size of the dynamic list for searching
                    metric="cosine"  # Distance metric (cosine, euclidean, dotProduct)
                )
            )
        ],
        profiles=[
            azsdim.VectorSearchProfile(
                name="default-profile",
                algorithm_configuration_name="default-algorithm"
            )
        ]
    )
    return vector_search

def get_exhaustive_KNN_vector_search(algorithm_name: str = "default-algorithm", 
                                     vector_search_profile_name: str = "default-profile", 
                                     metric: str = "cosine") -> azsdim.VectorSearch:
    """
    Get an exhaustive KNN vector search configuration.
    Args:
        algorithm_name: Name of the algorithm. Default is "default-algorithm".
        vector_search_profile_name: Name of the vector search profile. Default is "default-profile".
        metric: Distance metric (cosine, euclidean, dotProduct). Default is cosine.
    Returns:
        An exhaustive KNN vector search configuration
    """
    # Define vector search configuration
    vector_search = azsdim.VectorSearch(
        algorithms=[
            azsdim.ExhaustiveKnnAlgorithmConfiguration(
                name=algorithm_name,
                parameters=azsdim.ExhaustiveKnnParameters(
                    metric=metric
                )
            )
        ],
        profiles=[
            azsdim.VectorSearchProfile(
                name=vector_search_profile_name,
                algorithm_configuration_name=algorithm_name
            )
        ]
    )
    return vector_search

class SearchIndexerManager:
    """
    A manager for working with Azure AI Search indexers, data sources, and skillsets.
    """
    
    def __init__(self, search_service: SearchService):
        """
        Initialize a new SearchIndexerManager.
        
        Args:
            search_service: The SearchService to use
        """
        self.search_service = search_service
        self.indexer_client = search_service.get_indexer_client()
        
    # Data Source Connection methods
    def get_data_source_connections(self) -> List["DataSourceConnection"]:
        """
        Get all data source connections for this search service.
        
        Returns:
            List of DataSourceConnection objects
        """
        data_sources = self.indexer_client.get_data_source_connections()
        return [DataSourceConnection(self, ds) for ds in data_sources]
        
    def get_data_source_connection(self, name: str) -> Optional["DataSourceConnection"]:
        """
        Get a data source connection by name.
        
        Args:
            name: The name of the data source connection
            
        Returns:
            DataSourceConnection object or None if not found
        """
        try:
            data_source = self.indexer_client.get_data_source_connection(name)
            return DataSourceConnection(self, data_source)
        except Exception:
            return None
            
    def create_data_source_connection(self, name: str, type: str,
                                     connection_string: str,
                                     container: azsdim.SearchIndexerDataContainer) -> "DataSourceConnection":
        """
        Create a new data source connection.
        
        Args:
            name: The name of the data source connection
            type: The type of data source (e.g., "azureblob", "azuretable", "azuresql")
            connection_string: The connection string for the data source
            container: The container information
            
        Returns:
            DataSourceConnection object
        """
        data_source = azsdim.SearchIndexerDataSourceConnection(
            name=name,
            type=type,
            connection_string=connection_string,
            container=container
        )
        result = self.indexer_client.create_data_source_connection(data_source)
        return DataSourceConnection(self, result)
        
    # Indexer methods
    def get_indexers(self) -> List["Indexer"]:
        """
        Get all indexers for this search service.
        
        Returns:
            List of Indexer objects
        """
        indexers = self.indexer_client.get_indexers()
        return [Indexer(self, indexer) for indexer in indexers]
        
    def get_indexer(self, name: str) -> Optional["Indexer"]:
        """
        Get an indexer by name.
        
        Args:
            name: The name of the indexer
            
        Returns:
            Indexer object or None if not found
        """
        try:
            indexer = self.indexer_client.get_indexer(name)
            return Indexer(self, indexer)
        except Exception:
            return None
            
    def create_indexer(self, name: str, data_source_name: str,
                      target_index_name: str, skillset_name: Optional[str] = None,
                      schedule: Optional[azsdim.IndexingSchedule] = None,
                      field_mappings: Optional[List[azsdim.FieldMapping]] = None,
                      output_field_mappings: Optional[List[azsdim.OutputFieldMappingEntry]] = None,
                      parameters: Optional[azsdim.IndexingParameters] = None) -> "Indexer":
        """
        Create a new indexer.
        
        Args:
            name: The name of the indexer
            data_source_name: The name of the data source to use
            target_index_name: The name of the index to populate
            schedule: Optional indexing schedule
            parameters: Optional indexing parameters
            
        Returns:
            Indexer object
        """
        indexer = azsdim.SearchIndexer(
            name=name,
            data_source_name=data_source_name,
            target_index_name=target_index_name,
            skillset_name=skillset_name,
            schedule=schedule,
            parameters=parameters,
            field_mappings=field_mappings,
            output_field_mappings=output_field_mappings
        )
        result = self.indexer_client.create_indexer(indexer)
        return Indexer(self, result)
        
    # Skillset methods
    def get_skillsets(self) -> List["Skillset"]:
        """
        Get all skillsets for this search service.
        
        Returns:
            List of Skillset objects
        """
        skillsets = self.indexer_client.get_skillsets()
        return [Skillset(self, skillset) for skillset in skillsets]
        
    def get_skillset(self, name: str) -> Optional["Skillset"]:
        """
        Get a skillset by name.
        
        Args:
            name: The name of the skillset
            
        Returns:
            Skillset object or None if not found
        """
        try:
            skillset = self.indexer_client.get_skillset(name)
            return Skillset(self, skillset)
        except Exception:
            return None
            
    def create_skillset(self, name: str, skills: List[azsdim.SearchIndexerSkill],
                       description: Optional[str] = None) -> "Skillset":
        """
        Create a new skillset.
        
        Args:
            name: The name of the skillset
            skills: The skills to include in the skillset
            description: Optional description
            
        Returns:
            Skillset object
        """
        skillset = azsdim.SearchIndexerSkillset(
            name=name,
            skills=skills,
            description=description
        )
        result = self.indexer_client.create_skillset(skillset)
        return Skillset(self, result)

class DataSourceConnection:
    """
    Represents a data source connection in Azure AI Search.
    """
    
    def __init__(self, manager: SearchIndexerManager, data_source: azsdim.SearchIndexerDataSourceConnection):
        """
        Initialize a new DataSourceConnection.
        
        Args:
            manager: The SearchIndexerManager that created this object
            data_source: The underlying SearchIndexerDataSourceConnection
        """
        self.manager = manager
        self.data_source = data_source
        
    def get_name(self) -> str:
        """
        Get the name of this data source connection.
        
        Returns:
            The name of the data source connection
        """
        return self.data_source.name
        
    def update(self, connection_string: Optional[str] = None,
              container: Optional[azsdim.SearchIndexerDataContainer] = None) -> "DataSourceConnection":
        """
        Update this data source connection.
        
        Args:
            connection_string: Optional new connection string
            container: Optional new container
            
        Returns:
            Updated DataSourceConnection
        """
        if connection_string:
            self.data_source.connection_string = connection_string
        if container:
            self.data_source.container = container
            
        result = self.manager.indexer_client.create_or_update_data_source_connection(self.data_source)
        return DataSourceConnection(self.manager, result)
        
    def delete(self) -> None:
        """
        Delete this data source connection.
        """
        self.manager.indexer_client.delete_data_source_connection(self.data_source)

class Indexer:
    """
    Represents an indexer in Azure AI Search.
    """
    
    def __init__(self, manager: SearchIndexerManager, indexer: azsdim.SearchIndexer):
        """
        Initialize a new Indexer.
        
        Args:
            manager: The SearchIndexerManager that created this object
            indexer: The underlying SearchIndexer
        """
        self.manager = manager
        self.indexer = indexer
        
    def get_name(self) -> str:
        """
        Get the name of this indexer.
        
        Returns:
            The name of the indexer
        """
        return self.indexer.name
        
    def run(self) -> None:
        """
        Run this indexer.
        """
        self.manager.indexer_client.run_indexer(self.indexer.name)
        
    def reset(self) -> None:
        """
        Reset this indexer.
        """
        self.manager.indexer_client.reset_indexer(self.indexer.name)
        
    def get_status(self) -> azsdim.SearchIndexerStatus:
        """
        Get the status of this indexer.
        
        Returns:
            The status of the indexer
        """
        return self.manager.indexer_client.get_indexer_status(self.indexer.name)
        
    def update(self, schedule: Optional[azsdim.IndexingSchedule] = None,
              parameters: Optional[azsdim.IndexingParameters] = None) -> "Indexer":
        """
        Update this indexer.
        
        Args:
            schedule: Optional new indexing schedule
            parameters: Optional new indexing parameters
            
        Returns:
            Updated Indexer
        """
        if schedule:
            self.indexer.schedule = schedule
        if parameters:
            self.indexer.parameters = parameters
            
        result = self.manager.indexer_client.create_or_update_indexer(self.indexer)
        return Indexer(self.manager, result)
        
    def delete(self) -> None:
        """
        Delete this indexer.
        """
        self.manager.indexer_client.delete_indexer(self.indexer)
        
class Skillset:
    """
    Represents a skillset in Azure AI Search.
    """
    
    def __init__(self, manager: SearchIndexerManager, skillset: azsdim.SearchIndexerSkillset):
        """
        Initialize a new Skillset.
        
        Args:
            manager: The SearchIndexerManager that created this object
            skillset: The underlying SearchIndexerSkillset
        """
        self.manager = manager
        self.skillset = skillset
        
    def get_name(self) -> str:
        """
        Get the name of this skillset.
        
        Returns:
            The name of the skillset
        """
        return self.skillset.name
        
    def update(self, skills: Optional[List[azsdim.SearchIndexerSkill]] = None,
              description: Optional[str] = None) -> "Skillset":
        """
        Update this skillset.
        
        Args:
            skills: Optional new skills
            description: Optional new description
            
        Returns:
            Updated Skillset
        """
        if skills:
            self.skillset.skills = skills
        if description:
            self.skillset.description = description
            
        result = self.manager.indexer_client.create_or_update_skillset(self.skillset)
        return Skillset(self.manager, result)
        
    def delete(self) -> None:
        """
        Delete this skillset.
        """
        self.manager.indexer_client.delete_skillset(self.skillset)

class SearchIndex:
    index_name: str
    fields: List[azsdim.SearchField]
    vector_search: azsdim.VectorSearch
    search_service: SearchService
    azure_index: azsdim.SearchIndex

    def __init__(self, search_service: SearchService, index_name: str, fields: List[azsdim.SearchField], vector_search: Optional[azsdim.VectorSearch] = None):
        self.search_service = search_service
        self.index_name = index_name
        self.fields = fields
        self.vector_search = vector_search

        # SimpleField, SearchableField, ComplexField, are derived from SearchField
        index_definition = azsdim.SearchIndex(name=self.index_name, fields=fields, vector_search=vector_search)
        self.azure_index = self.search_service.get_index_client().create_or_update_index(index_definition)

    def get_search_client(self, index_name: Optional[str] = None ) -> azsd.SearchClient:
        if not index_name:
            index_name = self.index_name
        search_client = self.search_service.search_client            
        if search_client is None or search_client.index_name != index_name:
            search_client = azsd.SearchClient(
                endpoint=self.search_service.get_service_endpoint(),
                index_name=index_name,
                credential=self.search_service.get_credential()
            )
        return search_client  

    def extend_index_schema(self, new_fields: List[azsdim.SearchField] ) -> Optional[bool]:
        """
        Extend an Azure AI Search index schema with new fields
        
        Args:
            new_fields: List of new field definitions to add
            Indicative fields:
            azsdim.SimpleField(name="category", type=SearchFieldDataType.String, filterable=True, facetable=True),
            azsdim.SimpleField(name="publication_date", type=SearchFieldDataType.DateTimeOffset, filterable=True, sortable=True),
            azsdim.SearchableField(name="summary", type=SearchFieldDataType.String, analyzer_name="en.microsoft")
        """

        try:
            existing_field_names = [field.name for field in self.azure_index.fields]
            
            # Filter out any fields that already exist in the index
            fields_to_add = [field for field in new_fields if field.name not in existing_field_names]
            
            if not fields_to_add:
                print("No new fields to add - all specified fields already exist in the index")
                return True
                
            self.azure_index.fields.extend(fields_to_add)
            
            index_client: azsdi.SearchIndexClient = self.search_service.get_index_client()
            result: Optional[bool] = index_client.create_or_update_index(self.azure_index)
            print(f"Successfully extended index '{self.index_name}' with {len(fields_to_add)} new fields")
            
            # Return the updated index
            return result
            
        except Exception as e:
            print(f"Error extending index: {str(e)}")
            raise

    def process_data_in_batches(self, 
                                index_name: Optional[str] = None,
                                transaction: Callable[[List[Dict[str, Any]]], int] = None,
                                search_text: str = "*",
                                batch_size: int = 100 ) -> Tuple[int, int]:
        '''
        Process data in batches from an Azure AI Search index

        Args:
            index_name: Name of the index to process. If None, the current index is used. 
            transaction: Function to process a batch of documents. Gets a list of documents and returns the number of successful transactions. The transaction function could upload documents to the same or another index.
            batch_size: Number of documents to process in each batch
        
        Returns:
            Tuple of (succeeded_count, document_count)

        '''
        if index_name is None:
            index_name = self.index_name
        search_client = self.get_search_client(index_name)
        skip = 0
        document_count = 0
        succeeded_count = 0 
        
        while True:
            results = search_client.search( search_text=search_text, include_total_count=True, skip=skip, top=batch_size )
            
            # Get the total document count
            if document_count == 0:
                document_count = results.get_count()
                print(f"Retrieved {document_count} documents to process")
            
            documents_to_process: List[Dict[str, Any]] = []
            batch_documents: List[Dict[str, Any]] = list(results)
            if not batch_documents:
                break  # No more documents to process
            
            for doc in batch_documents:
                documents_to_process.append(doc)
            
            # Upload the batch to the target index
            if documents_to_process:
                transaction_result = transaction(documents_to_process)
                
                succeeded_count += transaction_result
                print(f"Processed batch: {transaction_result}/{len(documents_to_process)} documents (offset: {skip})")
            
            # Move to the next batch
            skip += batch_size
            
            # Check if we've processed all documents
            if skip >= document_count:
                break
        print(f"Successfully processed {succeeded_count}/{document_count} documents from  index '{index_name}'")
        return (succeeded_count, document_count)

    def upload_rows( self, documents: List[Dict[str, Any]], index_name: Optional[str] = None ) -> List[azsd._generated.models.IndexingResult]:
        """
        Upload documents to an Azure AI Search index
        
        Args:
            documents: List of documents to upload
            index_name: Name of the target index. If None, the current index is used.
        
        Returns:
            Number of successfully uploaded documents
        """
        try:
            if not index_name:
                index_name = self.index_name
            search_client = self.get_search_client(index_name)
            
            # Upload documents to the target index
            result = search_client.upload_documents(documents=documents)
            
            # Return the number of successfully uploaded documents
            return result
        except Exception as e:
            print(f"Error uploading new index rows: {str(e)}")
            raise e
    

    def copy_index_data(self, source_index_name: str, target_index_name: str, fields_to_copy: Optional[List[str]] = None, batch_size: int = 100) -> Tuple[int, int]:
        """
        Copy data from source index to target index, excluding the removed fields
        
        Args:
            source_index_name: Name of the source index. if None the current index is used
            target_index_name: Name of the target index. if None the current index is used
            fields_to_copy: List of field names to copy from source to target. If None, all fields are copied.
            batch_size: Number of documents to process in each batch
        Returns:
            Tuple of (succeeded_count, document_count)
        """
        if not source_index_name:
            source_index_name = self.index_name
        if not target_index_name:
            target_index_name = self.index_name
        if source_index_name == target_index_name:
            print("Source and target index names are the same. No action taken.")
            return (0, 0)
        
        target_client = self.get_search_client(target_index_name)
        
        def copy_and_upload_documents(documents: List[Dict[str, Any]]) -> int:
            documents_to_upload = []
            for doc in documents:
                # Create a new document with the selected fields
                new_doc = {key: value for key, value in doc.items() if not fields_to_copy or (key in fields_to_copy) }
                documents_to_upload.append(new_doc)            
            # Upload the batch to the target index
            succeeded = 0
            if documents_to_upload:
                result = target_client.upload_documents(documents=documents_to_upload)
                
                succeeded = sum(1 for r in result if r.succeeded)

            return succeeded

        result = self.process_data_in_batches(index_name = source_index_name, transaction=copy_and_upload_documents)
        return result

    def copy_index_structure(self, fields_to_copy: Optional[List[str]] = None, new_index_name: Optional[str] = None ) -> azsdim.SearchIndex:
        """
        Make a copy of an Azure AI Search index with a subset of fields.
        
        Args:
            fields_to_copy: List of field names to copy/replicate. If None, all fields are copied.
            new_index_name: Name for the new index (defaults to original_index_name + "_new")
        
        Returns:
            azsdim.SearchIndex: The created search index
        """

        try:
            original_index = self.azure_index
            if not new_index_name:
                new_index_name = f"{self.index_name}_new"

            new_fields = [field for field in original_index.fields if not fields_to_copy or field.name in fields_to_copy]
            
            if len(new_fields) == 0:
                print(f"None of the specified fields exist in index '{self.index_name}'")
                # Return a minimal SearchIndex to match the return type
                return azsdim.SearchIndex(name=self.index_name, fields=[])
            
            # Create a new index definition
            if hasattr(original_index, "semantic_settings"):
                semantic_settings = original_index.semantic_settings
            else:
                semantic_settings = None
                
            new_index = azsdim.SearchIndex(
                name=new_index_name,
                fields=new_fields,
                # Copy other index properties that might be important
                scoring_profiles=original_index.scoring_profiles,
                default_scoring_profile=original_index.default_scoring_profile,
                cors_options=original_index.cors_options,
                suggesters=original_index.suggesters,
                analyzers=original_index.analyzers,
                tokenizers=original_index.tokenizers,
                token_filters=original_index.token_filters,
                char_filters=original_index.char_filters,
                semantic_settings=semantic_settings,
                vector_search=original_index.vector_search
            )
            
            # Create the new index
            result = self.search_service.get_index_client().create_or_update_index(new_index)
            
            return result
            
        except Exception as e:
            print(f"Error copying index structure: {str(e)}")
            raise

    def perform_search(self, fields_to_select:List[str]=None, highlight_fields:str="chunk", highlight_pre_tag:str="<b>", highlight_post_tag:str="</b>", include_total_count:bool=True, filter_expression:Optional[str]=None, top:int=10,
                       query_text:Optional[str]=None, search_options:Optional[Dict[str, Any]]=None) -> azsd.SearchItemPaged[Dict[str, Any]]:
        search_options = {
            "include_total_count": include_total_count,
            "select": fields_to_select,
            "highlight_fields": highlight_fields,
            "highlight_pre_tag": highlight_pre_tag,
            "highlight_post_tag": highlight_post_tag
        }
        if filter_expression:
            search_options["filter"] = filter_expression
        if top:
            search_options["top"] = top
        search_client = self.get_search_client()
        results = search_client.search(query_text, **search_options)
        return results
    
    def get_adjacent_chunks(self, all_chunks: List[Dict[str, Any]]) -> Tuple[Dict[str, List[Dict[str, Any]]], Dict[Tuple[str, str], int]]:
        # Organize chunks by parent_id
        parent_chunks = defaultdict(list)
        for chunk in all_chunks:
            if 'parent_id' in chunk and 'chunk_id' in chunk:
                parent_chunks[chunk['parent_id']].append(chunk)
        
        all_chunks_by_parent: Dict[str, List[Dict[str, Any]]] = {}
        chunk_position_map: Dict[Tuple[str, str], int] = {}
        # Sort chunks within each parent document and create position maps
        for parent_id, chunks in parent_chunks.items():
            # Try to sort chunks within each parent document
            try:
                # First try to sort by chunk_id if it contains a number
                def extract_number(chunk: Dict[str, Any]) -> Union[int, str]:
                    chunk_id = chunk.get('chunk_id', '')
                    # Try to extract a number from the chunk_id
                    if '_' in chunk_id:
                        parts = chunk_id.split('_')
                        if parts[-1].isdigit():
                            return int(parts[-1])
                    if chunk_id.isdigit():
                        return int(chunk_id)
                    return chunk_id  # Fall back to string sorting
                
                sorted_chunks = sorted(chunks, key=extract_number)
            except:
                # If sorting fails, keep the original order
                sorted_chunks = chunks
            
            # Store the sorted chunks
            all_chunks_by_parent[parent_id] = sorted_chunks
            
            # Create a position map for quick lookup
            for i, chunk in enumerate(sorted_chunks):
                chunk_position_map[(parent_id, chunk['chunk_id'])] = i
            
        return all_chunks_by_parent, chunk_position_map

    def search_with_context_window(self,
                             query_text: str,
                             query_vector: List[float],
                             vector_fields: Optional[str] = None,
                             search_options: Optional[Dict[str, Any]] = None,
                             use_semantic_search: bool = False,
                             semantic_config_name: str = "default-semantic-config",
                             top: int = 10,
                             window_size: int = 3) -> List[Dict[str, Any]]:
        """
        Perform a search and retrieve a window of context chunks around each result.
        
        Args:
            query_text (str): The search query text
            window_size (int): Number of chunks to include before and after each result
            semantic_enabled (bool): Whether to enable semantic search
            top_k (int): Number of results to return
            vector_fields (str): comma separated(?) vector fields to search
            text_fields (list): List of text fields to search
            filter_condition (str): Optional OData filter condition
            
        Returns:
            list: Search results enriched with context windows
        """
        results = []
        try:
            # Get initial search results
            results = self.perform_hybrid_search(query_text=query_text, 
                                            query_vector=query_vector, 
                                            vector_fields=vector_fields, 
                                            use_semantic_search=use_semantic_search,
                                            top=top,
                                            semantic_config_name=semantic_config_name,
                                            search_options=search_options)
        except Exception as e:
            try:
                # Get initial search results without search options
                results = self.perform_hybrid_search(query_text=query_text, 
                                                query_vector=query_vector, 
                                                vector_fields=vector_fields, 
                                                use_semantic_search=use_semantic_search,
                                                top=top,
                                                semantic_config_name=semantic_config_name)
            except Exception as e:
                print(f"Error performing hybrid search: {e}")
                return []
            
        # Check if results are empty
        if not results:
            return []
        
        # If window size is specified, retrieve context chunks
        if window_size > 0:
            # Collect all parent_ids from the results
            parent_ids = set()
            for result in results:
                if 'parent_id' in result:
                    parent_ids.add(result['parent_id'])
            
            # Batch retrieve all chunks for all parent documents
            if parent_ids:
                # Create a filter to get all chunks from all parent documents in one query
                parent_filter = " or ".join([f"parent_id eq '{pid}'" for pid in parent_ids])
                
                # Retrieve all chunks (up to a reasonable limit)
                search_options: Dict[str, Any] = {
                    "include_total_count": True,
                    "select": "*"
                }
                all_chunks = list(self.perform_search("*", 
                                                    filter_expression=parent_filter,
                                                    top=1000,  # Adjust based on your needs
                                                    search_options=search_options))
                
                # Group chunks by parent_id
                all_chunks_by_parent, chunk_position_map = self.get_adjacent_chunks(all_chunks)
            
            # Enrich results with context windows
            enriched_results = []
            for result in results:
                parent_id = result.get('parent_id')
                chunk_id = result.get('chunk_id')
                
                # Initialize empty context window
                context_window: Dict[str, List[Dict[str, Any]]] = {
                    'before': [],
                    'after': []
                }
                
                if parent_id and chunk_id and parent_id in all_chunks_by_parent:
                    parent_chunks = all_chunks_by_parent[parent_id]
                    
                    # Find the position of this chunk
                    position = chunk_position_map.get((parent_id, chunk_id))
                    
                    if position is not None:
                        # Get previous chunks (up to window_size)
                        start_idx = max(0, position - window_size)
                        context_window['before'] = parent_chunks[start_idx:position]
                        
                        # Get next chunks (up to window_size)
                        end_idx = min(len(parent_chunks), position + window_size + 1)
                        context_window['after'] = parent_chunks[position+1:end_idx]
                
                enriched_result = {
                    'result': result,
                    'context_window': context_window
                }
                
                enriched_results.append(enriched_result)
            
            results_return = []
            for result in enriched_results:
                results_return.append(result['result'])
                for chunk in result['context_window']['before']:
                    results_return.append(chunk)
                for chunk in result['context_window']['after']:
                    results_return.append(chunk)

            return results_return
        else:
            # If no window size is specified, return the original results
            return results

    def perform_hybrid_search(self,
                              query_text: str,
                              query_vector: List[float],
                              vector_fields: Optional[str] = None,
                              search_options: Optional[Dict[str, Any]] = None,
                              use_semantic_search: bool = False,
                              top: int = 10,
                              semantic_config_name: str = "default-semantic-config",
                              display_fields: List[str] = None,
                              include_total_count: bool = True,
                              vectorized_query_kind: str = "vector",
                              exhaustive: bool = False,
                              k_nearest_neighbors_vector_search: int = 50,
                              query_answer: str = "extractive") -> List[Dict[str, Any]]:
        """
        Perform a hybrid search combining traditional keyword search with vector search.
        Args:
            query_text: The search query text
            query_vector: The vector representation of the query
            vector_fields: List of fields to perform vector search on
            search_options: Additional search options
            use_semantic_search: Whether to use a semantic search configuration
            top: The number of search results to retrieve
            semantic_config_name: The name of the semantic configuration to use
            display_fields: List of fields to display in the results
            include_total_count: Whether to include the total count of results in the response
            vectorized_query_kind: The kind of vector query being performed. Known values are: "vector" and "text".
            exhaustive: Whether to trigger an exhaustive k-nearest neighbor search across all vectors within the vector index
            k_nearest_neighbors_vector_search: Number of nearest neighbors to return as top hits
            query_answer: This parameter is only valid if the query type is 'semantic'. If set, the query returns answers extracted from key passages in the highest ranked documents
        Returns:
            A list of search results
        """
        # Default vector fields if not provided
        if vector_fields is None:
            vector_fields = "text_vector"
        
        # Create vectorized query
        vectorized_query = VectorizedQuery(vector=query_vector, 
                                           kind=vectorized_query_kind, 
                                           k_nearest_neighbors=k_nearest_neighbors_vector_search, 
                                           fields=vector_fields, 
                                           exhaustive=exhaustive)
        
        # Default search options
        default_options: Dict[str, Any] = {
            "search_text": query_text,  # Traditional keyword search
            "vector_queries": [vectorized_query],  # Vector search component
            "top": top,
            "select": display_fields,
            "include_total_count": include_total_count,
        }
        
        # Add semantic search if requested
        if use_semantic_search:
            default_options.update({
                "query_type": "semantic",
                "semantic_configuration_name": semantic_config_name,
                "query_caption": "extractive", 
                "query_answer": query_answer,
            })        
        
        # Update with any user-provided options
        if search_options:
            default_options.update(search_options)
        
        # Execute the search
        search_client = self.get_search_client()
        results = search_client.search(**default_options)
        
        # Process and return the results
        processed_results = []
        for result in results:
            processed_result = dict(result)
            processed_results.append(processed_result)
        
        return processed_results    
        
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.section import Section

class DocParsing:
    """
    Handles content from "docx" type files, to process them into documents for usage by "MultiProcessHandler" class. Utilizes the python-docx package.

    Parameters:
        doc_instance: python-docx Document object to be parsed
        openai_client: Azure OpenAI client for AI processing
        json_format: Template for the JSON structure
        domain: Domain category for the document
        sub_domain: Sub-domain category for the document
        model_name: Name of the AI model to use
        doc_name: Name of the document being processed (without extension)
    """
    doc_instance: DocumentObject
    openai_client: OpenAIClient
    json_format: Dict[str, Any]
    domain: str
    sub_domain: str
    model_name: str
    doc_name: str

    def __init__(self, doc_instance: DocumentObject, openai_client: OpenAIClient, json_format: Dict[str, Any], domain: str, sub_domain: str, model_name: str, doc_name: str):
        print(f"Initializing DocParsing for document: {doc_name}")
        self.doc_instance = doc_instance
        self.openai_client = openai_client
        self.json_format = json_format
        self.domain = domain
        self.sub_domain = sub_domain
        self.model_name = model_name
        self.doc_name = doc_name
    
    def _get_section_header_lines(self, section: Section) -> List[str]:
        """Helper function to extract text lines from a section's header.

        Args:
            section: The python-docx Section object for the header.
        
        Returns:
            lines: List of text lines from the section header.
        """
        try:
            if not section or not section.header:
                return []

            lines = []
            # Gather paragraph text from the header
            for paragraph in section.header.paragraphs:
                txt = paragraph.text.strip()
                if txt:
                    lines.append(txt)

            # Gather table cell text from the header (if any)
            for table in section.header.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        lines.append(" | ".join(row_cells)) # Join cell text for table lines
            return lines
        except Exception as e:
            print(f"Error extracting header lines for section: {e}")
            return []

    def _parse_header_lines(self, header_lines: List[str]) -> str:
        """Helper function to parse header lines to isolate the process title.

        Args:
            header_lines: List of text lines from the section header.
        
        Returns:
            The identified title segment in the provided lines.
        """
        if not header_lines:
            return "Metadata" # Default if no lines or only empty lines

        # Pattern for process numbers (e.g., 1., 1.1., 1.1.1.)
        number_pattern = re.compile(r'^\d+(\.\d+)*\.$')
        # Pattern for specific metadata lines to ignore (Example from doc_parsing.py)
        meta_patterns = [r'^Εκδ\.', r'Σελ\.'] # Add more patterns if needed

        potential_title = "Metadata" # Start with default

        for i, line in enumerate(header_lines):
            line_stripped = line.strip()
            if not line_stripped: continue # Skip empty lines

            # Skip known metadata lines
            if any(re.search(pattern, line_stripped) for pattern in meta_patterns):
                continue

            # Check if line matches "Number.\tTitle" format
            if "\t" in line_stripped:
                parts = line_stripped.split("\t", 1)
                potential_num = parts[0].strip()
                potential_title_part = parts[1].strip() if len(parts) > 1 else ""
                if number_pattern.match(potential_num) and potential_title_part:
                    return potential_title_part # Found title directly

            # Check if line is just a process number
            elif number_pattern.match(line_stripped):
                # Look for a non-metadata title in the *next* non-empty line
                if i + 1 < len(header_lines):
                    next_line_stripped = header_lines[i+1].strip()
                    if next_line_stripped and not any(re.search(pattern, next_line_stripped) for pattern in meta_patterns):
                         # Check if the next line looks like a title (heuristic: doesn't start with a number pattern)
                        if not number_pattern.match(next_line_stripped.split()[0] if next_line_stripped else ""):
                            return next_line_stripped # Found title on the next line

            # If the line is not metadata and not a number, consider it a potential title
            # This handles cases where title appears alone without a preceding number line
            elif potential_title == "Metadata": # Only take the first potential title
                 potential_title = line_stripped


        # If no specific pattern matched, return the first non-metadata line found, or "Metadata"
        return potential_title

    def _extract_header_info(self, section: Section) -> str:
        """Function to extracts the process title from a section header.
        
        Args:
            section: The python-docx Section object for the header.
        
        Returns:
            header_title: The header section's title.
        """
        try:
            lines = self._get_section_header_lines(section)
            header_title = self._parse_header_lines(lines)
            return header_title
        except Exception as e:
            print(f"Error extracting header info: {e}")
            return "Unknown Header" # Return a default on error

    def _iterate_block_items_with_section(self, doc: DocumentObject):
        """Function to iterate through the document's content blocks (paragraphs, tables).

        Args:
            doc: The python-docx Document object to parse.

        Yields:
            The section's index and the content block.
        """
        parent_elm = doc._element.body
        current_section_index = 0
        last_element_was_sectPr = False

        for child in parent_elm.iterchildren():
            if child.tag.endswith("p"):
                paragraph = Paragraph(child, doc)
                is_section_end_paragraph = bool(child.xpath("./w:pPr/w:sectPr"))
                if not is_section_end_paragraph:
                     yield current_section_index, paragraph
                if is_section_end_paragraph:
                    current_section_index += 1
                    last_element_was_sectPr = True
                else:
                    last_element_was_sectPr = False
            elif child.tag.endswith("tbl"):
                table = Table(child, doc)
                yield current_section_index, table
                last_element_was_sectPr = False
            elif child.tag.endswith('sectPr') and not last_element_was_sectPr:
                 current_section_index += 1

    def _extract_table_data(self, table: Table):
        """Extract text data from a table.
        
        Args: 
            Python-docx Table object.
        Returns:
            Joined cells with ' - '.
        """
        data = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                data.append(' - '.join(row_cells))
        return '\n'.join(data) # Join rows with newline

    def _is_single_process(self) -> Tuple[bool, str]:
        """Checks if the document contains a single process based on headers.
        
        Returns:
            True if single process, False if multi-process, and the title.
        """
        print("Checking document for single vs. multi-process structure...")
        section_headers = set()
        first_meaningful_header = None

        if not self.doc_instance.sections:
            print("Document has no sections.")
            return True, self.doc_name # Treat as single process with doc name as title

        for section_index, section in enumerate(self.doc_instance.sections):
            header_title = self._extract_header_info(section)
            if header_title and header_title != "Metadata" and header_title != "Unknown Header":
                section_headers.add(header_title)
                if first_meaningful_header is None:
                    first_meaningful_header = header_title # Store the first valid header found

        num_unique_headers = len(section_headers)
        print(f"Found {num_unique_headers} unique meaningful header(s): {section_headers if section_headers else 'None'}")

        if num_unique_headers <= 1: # Treat 0 or 1 unique headers as single process
            title = first_meaningful_header if first_meaningful_header else self.doc_name
            print(f"Document treated as single process with title: '{title}'")
            return True, title
        else:
            print("Document identified as multi-process.")
            return False, None # Title is None for multi-process

    def extract_data(self) -> Dict[Any, str]:
        """
        Extracts content from the document, handling single/multi-process structures.

        Returns:
            dict: Keys are formatted headers/process names, values are extracted content strings.
        """
        print("Extracting data based on document structure...")
        data_dict = {}
        is_single, process_title = self._is_single_process()

        if is_single:
            safe_title = re.sub(r'[\\/*?:"<>|]', '_', process_title) # Basic sanitization
            header_key = f"{safe_title}_single_process"
            print(f"Building content for single process: '{header_key}'")
            data_dict[header_key] = []
            for _, block in self._iterate_block_items_with_section(self.doc_instance):
                if isinstance(block, Paragraph):
                    text = block.text.strip()
                    if text: data_dict[header_key].append(text)
                elif isinstance(block, Table):
                    table_text = self._extract_table_data(block)
                    if table_text: data_dict[header_key].append(table_text)
        else:
            print("Building content for multi-process document...")
            last_section_index = -1
            current_header_key = None
            current_section_content = []

            for section_index, block in self._iterate_block_items_with_section(self.doc_instance):
                if section_index > last_section_index:
                    if current_header_key and current_section_content:
                         data_dict[current_header_key] = "\n".join(current_section_content) # Join previous section
                         print(f"Finalized content for section {last_section_index}: '{current_header_key}' ({len(current_section_content)} blocks)")

                    if section_index < len(self.doc_instance.sections):
                         header_title = self._extract_header_info(self.doc_instance.sections[section_index])
                         if not header_title or header_title == "Metadata":
                             header_title = f"Unknown_Section_{section_index}"
                         safe_header = re.sub(r'[\\/*?:"<>|]', '_', header_title)
                         current_header_key = f"{self.doc_name}_header_{safe_header}"
                         print(f"New Section {section_index}: Header='{header_title}', Key='{current_header_key}'")
                         if current_header_key not in data_dict:
                              data_dict[current_header_key] = []
                         current_section_content = [] # Reset buffer
                    else:
                         print(f"Warning: Block referenced section_index {section_index} > section count {len(self.doc_instance.sections)}. Using last header '{current_header_key}'.")

                    last_section_index = section_index

                block_text = ""
                if isinstance(block, Paragraph):
                    block_text = block.text.strip()
                elif isinstance(block, Table):
                    block_text = self._extract_table_data(block)

                if block_text and current_header_key:
                     # Append text directly to the list in the dictionary
                     if current_header_key in data_dict:
                         data_dict[current_header_key].append(block_text)
                     else:
                         # This case might happen if the first block has no preceding header info
                         print(f"Warning: No current header key for block, text ignored: '{block_text[:50]}...'")


            # Finalize the very last section after the loop
            if current_header_key and current_section_content:
                data_dict[current_header_key] = "\n".join(current_section_content)
                print(f"Finalized content for last section {last_section_index}: '{current_header_key}' ({len(current_section_content)} blocks)")


        # Join the collected content lines for each key into final strings
        final_data = {key: "\n".join(content_list).strip() for key, content_list in data_dict.items()}
        print(f"Data extraction complete. Found {len(final_data)} process/section block(s).")
        return final_data

    def update_json_with_ai(self, content_to_parse: str, process_identifier: str) -> str:
        """
        Uses AI to parse document content into the structured JSON format.

        Parameters:
            content_to_parse: The text content extracted for a specific process/section.
            process_identifier: The identifier (like header key) for this process/section.

        Returns:
            str: JSON string containing the parsed content, or None on failure.
        """
        print(f"Requesting AI to parse content for: '{process_identifier}' using model '{self.model_name}'...")
        format_str = json.dumps(self.json_format, indent=4, ensure_ascii=False)

        # Prompt emphasizing extraction, structure, and JSON-only output
        prompt = (
            "Parse the provided information about a specific process from the document and fill in the JSON structure below. "
            "Do not summarize, omit, or modify any details. Simply extract and organize the provided data into the corresponding fields of the JSON. "
            "There are more than one step and you have to include all of them.The step description has to be the whole text till the next step name"
            "Ensure every relevant detail is included without altering the content. "
            "The JSON format should follow this structure and include all fields, even if some of them are not present in the content (leave them empty or null if necessary):\n"
            f"{format_str}\n\n"
            "To make it clear the content you generate will be ONLY THE CONTENT of a json no \\n nothing.The first character {{ and the last character should be }}" # Escaped braces
            "Your response should be ONLY a JSON file content ready to be stored as json without other processing, with the exact format as shown above."
         )

        try:
            if not self.openai_client:
                 print("Error: Azure OpenAI client is not initialized.")
                 return None

            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": f"Document Source Name: {self.doc_name}\nProcess/Section Identifier: {process_identifier}\n\nContent to Parse:\n---\n{content_to_parse}\n---"}
            ]

            output_llm = self.openai_client.generate_chat_completion(
                model=self.model_name,
                messages=messages,
                temperature=0,
                max_tokens=None,
                response_format={"type": "json_object"} # Request JSON output
            )

            ai_response_content = output_llm["content"]
            print(f"AI response received ({len(ai_response_content)} chars).")

            # Basic validation: Check if it looks like JSON
            if ai_response_content and ai_response_content.strip().startswith("{") and ai_response_content.strip().endswith("}"):
                return ai_response_content.strip()
            else:
                print(f"Warning: AI response does not look like valid JSON: {ai_response_content[:100]}...")
                # Attempt to extract JSON if response_format failed or wasn't respected
                match = re.search(r'\{.*\}', ai_response_content, re.DOTALL)
                if match:
                    print("Extracted potential JSON from response.")
                    return match.group(0)
                else:
                    print("Error: Could not extract JSON object from AI response.")
                    return None

        except Exception as e:
            print(f"Error during AI call for '{process_identifier}': {e}")
            return None

    def _process_doc(self, ai_json_string: str) -> Dict[str, Any]:
        """
        Processes the AI response string, validates JSON and adds metadata.

        Parameters:
            ai_json_string: Raw JSON string from the AI.
        """
        try:
            # Parse the AI's JSON string into a Python dictionary
            json_data = json.loads(ai_json_string)

            # Create a new ordered dictionary to control the final structure
            ordered_data = {}
            ordered_data["doc_name"] = self.doc_name
            ordered_data["process_name"] = json_data.get("process_name", "Unknown Process Name")
            ordered_data["domain"] = self.domain
            ordered_data["subdomain"] = self.sub_domain

            # Add the rest of the fields from the AI response
            for key, value in json_data.items():
                if key not in ordered_data:
                    ordered_data[key] = value
            return ordered_data

        except json.JSONDecodeError as e:
            print(f"Error decoding JSON from AI response: {e}")

    def doc_to_json(self) -> List[Dict[str, Any]]:
        """
        Main method: Converts the python-docx object into a list of dictionaries to be indexed.

        Returns:
            The final list of dictionaries from the python-docx object.
        """
        print(f"Starting document-to-JSON conversion for '{self.doc_name}'...")

        extracted_data_dict = self.extract_data()
        
        ordered_data = []
        for process_key, content in extracted_data_dict.items():
            print(f"\n   Processing section/process key: '{process_key}'")
            if "Metadata" in process_key or "unknown" in process_key.lower() or not content.strip():
                print("Skipping metadata section or empty content.")
                continue

            filename_base = process_key.split('_header_')[-1] if '_header_' in process_key else process_key.replace('_single_process', '')
            safe_filename_base = re.sub(r'[\\/*?:"<>|]', '_', filename_base)
            max_len = 100 # Max filename length
            safe_filename_base = safe_filename_base[:max_len] if len(safe_filename_base) > max_len else safe_filename_base

            ai_json_result = self.update_json_with_ai(content, process_key)

            if ai_json_result:
                ordered_data.append(self._process_doc(ai_json_result))
            else:
                print(f"AI parsing failed for '{process_key}'.")
                print(f"{self.doc_name} - Section '{process_key}' - AI Parsing Failed")


        print(f"\nDocument-to-list conversion completed for '{self.doc_name}'")
        return ordered_data

import hashlib

class MultiProcessHandler:
    """
    Sets up the handler to process multiple documents and upload them to Azure Search
    using the provided clients.
    
    Parameters:
        dict_list: List of dictionaries to process
        index_client_core: Azure SearchIndex client for the core index
        index_client_detail: Azure SearchIndex client for the detailed index
        openai_client: Azure OpenAI client for generating embeddings
    """
    dict_list: List[Dict]
    index_client_core: SearchIndex
    index_client_detail: SearchIndex
    openai_client: OpenAIClient

    def __init__(self, dict_list: List[str], index_client_core: SearchIndex, index_client_detail: SearchIndex, openai_client: OpenAIClient):
        self.dict_list = dict_list
        self.index_client_core = index_client_core
        self.index_client_detail = index_client_detail
        self.openai_client = openai_client
    
    def generate_process_id(self, process_name: str, non_llm_summary: str, doc_name: str) -> int:
        """
        Generate a unique integer ID for the process based on its name and short description.
        
        Creates a SHA-256 hash of the combined process name and description,
        then converts it to an integer ID.
        
        Parameters:
            process_name: Name of the process
            non_llm_summary: A non-LLM summary of the process
            
        Returns:
            String representation of the process ID derived from the hash
        """
        print(f"Generating Process ID")
        print(f"Process Name: {process_name}")
        print(f"A non-llm summary: {non_llm_summary}")
        print(f"Document Name: {doc_name}")
        
        content_to_hash = f"{process_name}-{non_llm_summary}--{doc_name}"
        hashed_content = hashlib.sha256(content_to_hash.encode('utf-8')).hexdigest()
        
        # Convert the hex string to an integer and return only the first 10 digits of the integer
        full_id = int(hashed_content, 16)
        process_id = str(full_id)
        
        print(f"Generated Process ID: {process_id}")
        return process_id

    def generate_step_id(self, process_name: str, step_name: str, step_content: str) -> int:
        """
        Generate a unique integer ID for the step.
        
        Creates a SHA-256 hash of the combined process name, step name, and step content,
        then converts it to an integer ID.
        
        Parameters:
            process_name: Name of the parent process
            step_name: Name of the step
            step_content: Content/description of the step
            
        Returns:
            String representation of the step ID derived from the hash
        """
        print(f"Generating Step ID")
        print(f"Process Name: {process_name}")
        print(f"Step Name: {step_name}")
        
        content_to_hash = f"{process_name}-{step_name}-{step_content}"
        hashed_content = hashlib.sha256(content_to_hash.encode('utf-8')).hexdigest()
        
        # Convert the hex string to an integer and return only the first 10 digits of the integer
        full_id = int(hashed_content, 16)
        step_id = str(full_id)
        
        print(f"Generated Step ID: {step_id}")
        return step_id

    def prepare_core_df_record(self, provided_dict:Dict) -> Tuple[str, Dict]:
        """
        Prepare record for core_df index.
        
        Creates a dictionary containing the main process information
        and a non-LLM summary combining various process attributes.
        
        Parameters:
            process_id: The unique ID for this process
            provided_dict: The process' dictionary
            
        Returns:
            Dictionary containing the core process information formatted for database storage
        """
        print("Preparing Core DataFrame Record")
        
        # Prepare steps information
        steps_info = []
        for step in provided_dict.get('steps', []):
            step_text = f"Βήμα {step['step_number']} {step['step_name']}"
            steps_info.append(step_text)

        # Prepare summary
        summary_parts = [
            "Εισαγωγή:", provided_dict.get('introduction', ''),
            "Σύντομη περιγραφή:", provided_dict.get('short_description', ''),
            "Αναλυτικά βήματα:", "\n".join(steps_info),
            "Οικογένεια προιόντων:", ", ".join(provided_dict.get('related_products', [])),
            "Έγγραφα αναφοράς:", ", ".join(provided_dict.get('reference_documents', []))
        ]
        non_llm_summary = "\n\n".join(summary_parts)

        process_name = provided_dict.get('process_name', '')
        doc_name = provided_dict.get('doc_name', '')
        process_id = self.generate_process_id(process_name, non_llm_summary, doc_name)

        # Prepare core record
        core_record = {
            'process_id': process_id,
            'process_name': process_name,
            'doc_name': doc_name,
            'domain': provided_dict.get('domain', ''),
            'sub_domain': provided_dict.get('subdomain', ''),
            'functional_area': provided_dict.get('functional_area', ''),
            'functional_subarea': provided_dict.get('functional_subarea', ''),
            'process_group': provided_dict.get('process_group', ''),
            'process_subgroup': provided_dict.get('process_subgroup', ''),
            'systems_manuals_used': ', '.join(provided_dict.get('systems_manuals_used', [])),
            'forms_documents': ', '.join(provided_dict.get('forms_documents', [])),
            'reference_documents': ', '.join(provided_dict.get('reference_documents', [])),
            'related_products': ', '.join(provided_dict.get('related_products', [])),
            'additional_information': provided_dict.get('additional_information', ''),
            'process_added_at': datetime.now(timezone.utc),
            'non_llm_summary': non_llm_summary.strip()
        }

        print("Core DataFrame Record prepared successfully")
        return process_id, core_record

    def prepare_detailed_df_records(self, process_id: int, provided_dict:Dict) -> List[Dict]:
        """
        Prepare records for detailed_df index.
        
        Creates a list of dictionaries, each containing information about a step
        in the process, including an introduction record (step 0) and all regular steps.
        
        Parameters:
            process_id: The unique ID for the parent process
            provided_dict: The process' dictionary
            
        Returns:
            List of dictionaries containing detailed step information formatted for database storage
        """
        print("Preparing Detailed DataFrame Records")
        detailed_records = []

        # Generate Process ID
        process_name = provided_dict.get('process_name', '')

        # Add Introduction (step 0)
        intro_content = (
            f"Εισαγωγή:\n{provided_dict.get('introduction', '')}\n\n"
            f"Σύντομη περιγραφή:\n{provided_dict.get('short_description', '')}\n\n"
            f"Οικογένεια προιόντων:\n{', '.join(provided_dict.get('related_products', []))}\n\n"
            f"Έγγραφα αναφοράς:\n{', '.join(provided_dict.get('reference_documents', []))}"
        )
        
        intro_record = {
            'id': self.generate_step_id(process_name, "Εισαγωγή", intro_content),
            'process_id': process_id,
            'step_number': 0,
            'step_name': "Εισαγωγή",
            'step_content': intro_content.strip(),
            'documents_used': None,
            'systems_used': None
        }
        detailed_records.append(intro_record)

        # Add regular steps
        print(f"Total Steps: {len(provided_dict.get('steps', []))}")
        for step in provided_dict.get('steps', []):
            step_content = step.get('step_description', '')
            record = {
                'id': self.generate_step_id(process_name, step['step_name'], step_content),
                'process_id': process_id,
                'step_number': int(step['step_number']),
                'step_name': step['step_name'],
                'step_content': step_content,
                'documents_used': ', '.join(step.get('documents_used', [])),
                'systems_used': ', '.join(step.get('systems_used', []))
            }
            detailed_records.append(record)
            print(f"Step {record['step_number']}: {record['step_name']}")

        print("Detailed DataFrame Records prepared successfully")
        return detailed_records

    def prepare_for_upload(self, provided_dict) -> List[Dict]:
        """
        Prepares all records for upload.
        
        Coordinates the generation of process IDs and the preparation
        of both core and detailed records for upload.

        Parameters:
            provided_dict: The process' dictionary
        
        Returns:
            Tuple containing the core record dictionary and a list of detailed record dictionaries
        """
        print("Preparing records for upload")
        
        # Prepare core record
        process_id, core_record = self.prepare_core_df_record(provided_dict)

        # Prepare detailed records
        detailed_records = self.prepare_detailed_df_records(process_id, provided_dict)

        print("Upload preparation completed")
        # Combine the core record with the detailed records
        return core_record, detailed_records
    
    def process_documents(self) -> List[Dict]:
        """
        Processes multiple documents and returns a list of processed records for each document.
        
        Iterates through each document and uses the prepare_for_upload method
        to prepare the core and detailed records for upload.
        
        Returns:
            List of dictionaries, each containing 'core' and 'detailed' records for a document
        """
        all_records = []

        for i in self.dict_list:
            try:
                core_record, detailed_records = self.prepare_for_upload(i)
                all_records.append({
                    'core': core_record,
                    'detailed': detailed_records
                })
            except Exception as e:
                print(f"Error processing {self.dict_list}: {e}")

        return all_records
    
    def upload_to_azure_index(self, all_records: List[Dict], core_index_name: Optional[str] = None, detail_index_name: Optional[str] = None) -> None:
        """
        Uploads the processed records to Azure Search indexes.
        
        Generates embeddings for text fields and uploads the enriched records to Azure Search.
        For core records, creates embeddings for the summary.
        For detailed records, creates embeddings for both step names and step content.
        
        Parameters:
            all_records: List of processed records (each containing 'core' and 'detailed' keys)
            core_index_name: Name of the Azure Search index for core records
            detailed_index_name: Name of the Azure Search index for detailed records
            
        Note:
            This method ensures all IDs are converted to strings before upload
            and handles any errors that occur during the upload process.
            
        Results:
            Records are uploaded to Azure Search.
        """
        
        index_client_core = core_index_name or self.index_client_core
        index_client_detail = detail_index_name or self.index_client_detail

        openai_client = self.openai_client

        def get_embeddings_list(openai_client: OpenAIClient, texts: List[str], model: str = 'text-embedding-3-large') -> List[List[float]]:
            """
            Retrieve embeddings for a list of given texts.
            Handles text truncation if content exceeds model's token limit.
            """

            embeddings = []
            for text in texts:
                if text:
                    try:
                        embedding = openai_client.generate_embeddings(text=text, model=model)
                        embeddings.append(embedding)
                    except Exception as e:
                        # Check if error is related to token limit
                        error_str = str(e)
                        if "maximum context length" in error_str and "tokens" in error_str:
                            try:
                                # Try to extract the requested tokens from error message
                                import re
                                match = re.search(r'requested (\d+) tokens', error_str)
                                requested_tokens = int(match.group(1)) if match else 10000
                                max_tokens = 8192  # Default max tokens for embedding models
                                
                                if "maximum context length is" in error_str:
                                    max_match = re.search(r'maximum context length is (\d+)', error_str)
                                    if max_match:
                                        max_tokens = int(max_match.group(1))
                                
                                # Calculate ratio to truncate text
                                ratio = max_tokens / requested_tokens * 0.9  # 10% safety margin
                                truncated_length = int(len(text) * ratio)
                                truncated_text = text[:truncated_length]
                                print(f"Truncating text from {len(text)} chars to {len(truncated_text)} chars to fit token limit")
                                
                                # Try again with truncated text
                                embedding = openai_client.generate_embeddings(text=truncated_text, model=model)
                                embeddings.append(embedding)
                            except Exception as inner_e:
                                print(f"Error after truncation attempt: {inner_e}")
                                embeddings.append([])
                        else:
                            print(f"Error generating embeddings: {e}")
                            embeddings.append([])
                else:
                    embeddings.append([])
            return embeddings
        
        try:
            for record in all_records:
                # For the core record, generate an embedding for 'non_llm_summary' if it exists.
                if 'non_llm_summary' in record['core']:
                    summary_text = record['core']['non_llm_summary']
                    embeddings = get_embeddings_list(openai_client, [summary_text])
                    if embeddings and len(embeddings) > 0:
                        # Assign the embedding vector (list of numbers) directly
                        record['core']['embedding_summary'] = embeddings[0]
                
                # For each step record in the detailed part, generate embeddings for step_name and step_content.
                for step in record['detailed']:
                    # Ensure the step id is a string
                    if 'id' in step:
                        step['id'] = str(step['id'])
                    if 'step_name' in step:
                        name_embeddings = get_embeddings_list(openai_client, [step['step_name']])
                        if name_embeddings and len(name_embeddings) > 0:
                            step['embedding_title'] = name_embeddings[0]
                    if 'step_content' in step:
                        content_embeddings = get_embeddings_list(openai_client, [step['step_content']])
                        if content_embeddings and len(content_embeddings) > 0:
                            step['embedding_content'] = content_embeddings[0]
                
                record['core']['process_id'] = str(record['core']['process_id'])
                for i in record['detailed']:
                    i['id'] = str(i['id'])

                
                # Now upload the records to the respective Azure Search indexes.
                response_core = index_client_core.upload_rows(documents=[record['core']])
                response_detail = index_client_detail.upload_rows(documents=record['detailed'])
                print(f"Successfully uploaded records for {record['core'].get('process_name', 'Unknown')}")
        except Exception as e:
            print(f"Error uploading records: {e}")

