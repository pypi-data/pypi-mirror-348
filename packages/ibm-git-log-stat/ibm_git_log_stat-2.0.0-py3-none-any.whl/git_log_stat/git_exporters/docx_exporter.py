from collections import defaultdict
from logging import Logger
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from git_log_stat.git_exporters.base_exporter import BaseExporter
import textwrap


def generate_commit_summary_docx(log: Logger, commit_logs: str | list[str], pr_logs: str | list[str],
                                 output_path="commit_summary.docx", skip_summary: bool = False):
    # Normalize input
    if isinstance(commit_logs, list):
        commit_logs = "\n".join(commit_logs)
    if isinstance(pr_logs, list):
        pr_logs = "\n".join(pr_logs)

    grouped_commits = defaultdict(list)
    combined_messages = []

    for line in commit_logs.split("\n"):
        parts = [p.strip() for p in line.split('|')]
        if len(parts) >= 4:
            date, author, commit_id, *message_parts = parts
            message = " | ".join(message_parts)
            grouped_commits[author].append((date, message))
            combined_messages.append(f"{author} - {message}")

    summary = ""
    if not skip_summary:
        try:
            from git_log_stat.git_repo.git_nlp import init_summarizer
            summarizer = init_summarizer()
        except ImportError:
            summarizer = None

        log.info("🔄 Summarizing commit logs...")
        full_text = ". ".join(combined_messages)
        full_text = " ".join(full_text.split()[:1024])
        summary = summarizer(full_text, max_length=120, min_length=40, do_sample=False)[0]['summary_text']

    # Start DOCX
    doc = Document()

    # Title
    title = doc.add_heading("Commit Summary Report", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Spacer

    # Summary
    if summary:
        doc.add_heading("📝 Natural Language Summary", level=1)
        for line in textwrap.wrap(summary, width=100):
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)
        doc.add_paragraph()  # Spacer

    # Commits
    doc.add_heading("📜 Commits by Author", level=1)

    for author, commits in grouped_commits.items():
        doc.add_heading(f"👤 {author}", level=2)
        for date, message in commits:
            is_pr = "pull request" in message.lower() or "merge" in message.lower()
            formatted = f"🔀 {date} | {message}" if is_pr else f"{date} | {message}"
            p = doc.add_paragraph()
            run = p.add_run(formatted)
            run.font.size = Pt(10)
            if is_pr:
                run.bold = True
        doc.add_paragraph()  # Spacer between authors

    # PR Section (Optional, explicit)
    if pr_logs:
        doc.add_heading("📦 Pull Requests", level=1)
        for pr in pr_logs.strip().splitlines():
            for line in textwrap.wrap(pr.strip(), width=100):
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(10)

    # Save file
    doc.save(output_path)
    log.info(f"✅ DOCX saved to {output_path}")


class DocxExporter(BaseExporter):
    def export(self, output_file_name, commit_output, pr_output=None, skip_summary=False):
        self.log.info("Generating DOCX %s", output_file_name)
        generate_commit_summary_docx(
            self.log,
            commit_logs=commit_output,
            pr_logs=pr_output,
            output_path=output_file_name,
            skip_summary=skip_summary
        )
