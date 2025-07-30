"""
Utility functions for resume_ats plugin.
"""

import logging
import re
from typing import Dict, List, Set, Tuple

import nltk
from pypdf import PdfReader  # Updated from PyPDF2 to pypdf
from docx import Document
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Initialize logging
logger = logging.getLogger(__name__)

# Download NLTK resources if needed
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Get English stopwords
STOP_WORDS = set(stopwords.words('english'))


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    try:
        reader = PdfReader(pdf_path)
        text = ""
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF {pdf_path}: {e}")
        return ""


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        doc = Document(docx_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {docx_path}: {e}")
        return ""


def normalize_text(text: str) -> str:
    """
    Normalize text by removing extra whitespace, etc.
    
    Args:
        text: Text to normalize
        
    Returns:
        Normalized text
    """
    # First, preserve double newlines by replacing them with a placeholder
    text = re.sub(r'\n\s*\n', 'DOUBLE_NEWLINE_PLACEHOLDER', text)
    
    # Replace multiple whitespace with single space
    text = re.sub(r'\s+', ' ', text)
    
    # Restore double newlines
    text = text.replace('DOUBLE_NEWLINE_PLACEHOLDER', '\n\n')
    
    # Remove leading/trailing whitespace
    text = text.strip()
    
    return text


def get_keywords_from_job_description(job_description: str, max_keywords: int = 30) -> List[str]:
    """
    Extract key terms from job description.
    
    Args:
        job_description: Job description text
        max_keywords: Maximum number of keywords to extract
        
    Returns:
        List of extracted keywords
    """
    # Normalize job description
    job_text = normalize_text(job_description.lower())
    
    # Tokenize
    tokens = word_tokenize(job_text)
    
    # Filter stopwords and non-alphabetic tokens
    filtered_tokens = [w for w in tokens if w.isalpha() and len(w) > 3 and w not in STOP_WORDS]
    
    # Extract n-grams (terms that appear together)
    vectorizer = CountVectorizer(ngram_range=(1, 2), stop_words="english", max_features=max_keywords)
    X = vectorizer.fit_transform([job_text])
    
    # Get the most frequent terms
    feature_names = vectorizer.get_feature_names_out()
    frequencies = X.toarray()[0]
    
    # Sort by frequency
    keywords = [(feature_names[i], frequencies[i]) for i in range(len(feature_names))]
    keywords.sort(key=lambda x: x[1], reverse=True)
    
    # Return the top keywords
    top_keywords = [k[0] for k in keywords[:max_keywords]]
    
    # Add specific important technologies mentioned in the job description
    tech_keywords = ["python", "django", "flask", "react", "javascript", "aws", "docker", "kubernetes", "ci/cd", "git"]
    for keyword in tech_keywords:
        if keyword in job_text and keyword not in top_keywords:
            top_keywords.append(keyword)
    
    # Add important skill terms that might not be frequent
    skill_patterns = [
        r"(?:proficient|experience|familiar|knowledge|expertise)\s+(?:in|with)\s+([A-Za-z0-9_\+\#\-\.]+)",
        r"([A-Za-z0-9_\+\#]+)\s+(?:skills|proficiency)",
        r"(?:programming|language|framework|database)\s+(?:in|with)?\s*([A-Za-z0-9_\+\#]+)"
    ]
    
    for pattern in skill_patterns:
        for match in re.finditer(pattern, job_text):
            skill = match.group(1).lower().strip()
            if skill and skill not in top_keywords and skill not in STOP_WORDS:
                top_keywords.append(skill)
    
    # Deduplicate
    return list(dict.fromkeys(top_keywords))


def calculate_similarity_score(text1: str, text2: str) -> float:
    """
    Calculate cosine similarity between two texts.
    
    Args:
        text1: First text
        text2: Second text
        
    Returns:
        Similarity score between 0 and 1
    """
    vectorizer = CountVectorizer().fit_transform([text1, text2])
    vectors = vectorizer.toarray()
    
    if vectors.shape[0] < 2:
        return 0.0
    
    return cosine_similarity(vectors)[0, 1]


def extract_education_details(education_text: str) -> List[Dict]:
    """
    Extract structured education information.
    
    Args:
        education_text: Text from education section
        
    Returns:
        List of dictionaries with education details
    """
    education_entries = []
    degree_pattern = r'(?i)(bachelor|master|ph\.?d|doctorate|associate|b\.?s|m\.?s|b\.?a|m\.?a)'
    
    # Split by newline or multiple spaces to identify separate entries
    entries = re.split(r'\n+|\s{2,}', education_text)
    
    for entry in entries:
        if not entry.strip():
            continue
            
        result = {}
        
        # Look for degree
        degree_match = re.search(degree_pattern, entry)
        if degree_match:
            result['degree'] = degree_match.group(0)
        
        # Look for university/institution
        university_match = re.search(r'(?i)(university|college|institute|school) (?:of )?([A-Za-z\s]+)', entry)
        if university_match:
            result['institution'] = university_match.group(0)
        
        # Look for graduation year
        year_match = re.search(r'(?:19|20)\d{2}', entry)
        if year_match:
            result['year'] = year_match.group(0)
        
        # Add the entry if we found at least some information
        if result:
            result['text'] = entry.strip()
            education_entries.append(result)
    
    return education_entries


def extract_skills(text: str) -> List[str]:
    """
    Extract likely skills from text.
    
    Args:
        text: Resume text
        
    Returns:
        List of identified skills
    """
    # Common programming languages, tools, frameworks
    common_tech_skills = {
        'python', 'java', 'javascript', 'js', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin',
        'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring', 'rails',
        'sql', 'nosql', 'mongodb', 'mysql', 'postgresql', 'oracle', 'firebase',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins',
        'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence',
        'html', 'css', 'sass', 'less', 'bootstrap', 'tailwind',
        'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy'
    }
    
    # Common soft skills with regex to match variations
    soft_skills_patterns = [
        r'\bcommunicat(?:ion|e|ing|ed)\b',
        r'\bteam(?:work|player)\b',
        r'\bproblem.{1,5}solv(?:ing|e|ed)\b',
        r'\btime.{1,5}manage(?:ment)\b',
        r'\blead(?:ership|ing|er)\b',
        r'\banalytic(?:s|al)\b',
        r'\borganiz(?:e|ed|ation|ing)\b',
        r'\bcritical.{1,5}think(?:ing|er)\b',
        r'\bcreativ(?:e|ity)\b',
    ]
    
    found_skills = set()
    
    # Extract exact tech skills
    words = re.findall(r'\b[A-Za-z0-9#\+\._-]+\b', text.lower())
    for word in words:
        if word in common_tech_skills:
            found_skills.add(word)
    
    # Extract soft skills using patterns
    for pattern in soft_skills_patterns:
        matches = re.findall(pattern, text.lower())
        if matches:
            # Get the actual matched text from the original
            orig_matches = re.finditer(pattern, text.lower())
            for match in orig_matches:
                found_skills.add(match.group(0))
    
    # Deduplicate and sort
    return sorted(list(found_skills))