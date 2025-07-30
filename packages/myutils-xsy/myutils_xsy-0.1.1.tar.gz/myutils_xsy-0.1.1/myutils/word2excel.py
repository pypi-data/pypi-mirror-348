'''
Author: Shengyi Xu 54436848+xushengyichn@users.noreply.github.com
Date: 2025-05-12 15:07:11
LastEditors: Shengyi Xu 54436848+xushengyichn@users.noreply.github.com
LastEditTime: 2025-05-12 22:46:14
FilePath: /myutils/myutils/word2excel.py
Description: 

Copyright (c) 2025 by ${git_name_email}, All Rights Reserved. 
'''


# %%
import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import pandas as pd

def extract_title_body_tables_flat(doc_path):
    doc = Document(doc_path)

    # 提取段落
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    title = paragraphs[0] if paragraphs else ""
    body = "\n".join(paragraphs[1:]) if len(paragraphs) > 1 else ""

    # 提取表格（单元格级别）
    flat_table_data = {}
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                key = f"表{table_idx+1}_R{row_idx+1}_C{col_idx+1}"
                flat_table_data[key] = cell.text.strip()
    return title, body, flat_table_data

def confirm_dialog(title, body):
    root = tk.Tk()
    root.title("确认标题和正文")

    tk.Label(root, text="标题：", font=('Arial', 10, 'bold')).pack(anchor='w')
    title_box = tk.Text(root, height=2, width=80)
    title_box.pack()
    title_box.insert(tk.END, title)

    tk.Label(root, text="正文：", font=('Arial', 10, 'bold')).pack(anchor='w')
    body_box = tk.Text(root, height=20, width=80)
    body_box.pack()
    body_box.insert(tk.END, body)

    def on_confirm():
        root.destroy()

    def on_cancel():
        root.destroy()
        raise Exception("用户取消操作")

    tk.Button(root, text="确认写入", command=on_confirm).pack(side=tk.LEFT, padx=10, pady=5)
    tk.Button(root, text="取消", command=on_cancel).pack(side=tk.RIGHT, padx=10, pady=5)
    root.mainloop()

def batch_process(folder_path):
    results = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            full_path = os.path.join(folder_path, filename)
            title, body, flat_tables = extract_title_body_tables_flat(full_path)
            try:
                confirm_dialog(title, body)
                entry = {
                    "文件名": filename,
                    "标题": title,
                    "正文": body
                }
                entry.update(flat_tables)  # 加入所有单元格列
                results.append(entry)
            except:
                print(f"跳过文件：{filename}")
    return results

def main():
    folder_path = filedialog.askdirectory(title="选择Word文件所在文件夹")
    if not folder_path:
        return

    data = batch_process(folder_path)
    if data:
        df = pd.DataFrame(data)
        save_path = filedialog.asksaveasfilename(
            title="保存为Excel文件",
            defaultextension=".xlsx",
            filetypes=[("Excel 文件", "*.xlsx")]
        )
        if save_path:
            df.to_excel(save_path, index=False)
            messagebox.showinfo("完成", f"已保存为：{save_path}")
    else:
        messagebox.showwarning("无内容", "没有处理任何文件")

if __name__ == "__main__":
    main()

# %%
